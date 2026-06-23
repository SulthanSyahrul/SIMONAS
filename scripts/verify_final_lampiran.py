import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    print("=== Sequential Paragraphs from index 330 ===")
    for idx in range(330, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        drawings = p._element.xpath('.//w:drawing')
        picts = p._element.xpath('.//w:pict')
        images_count = len(drawings) + len(picts)
        text_preview = p.text.strip().replace('\n', ' [NL] ')[:80]
        print(f"P {idx:03d}: '{text_preview}' (images={images_count})")

if __name__ == '__main__':
    main()

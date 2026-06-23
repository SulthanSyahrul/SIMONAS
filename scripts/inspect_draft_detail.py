import docx

doc = docx.Document('Draft Template Naskah TA 2025.docx')
with open('scripts/inspect_draft_detail.txt', 'w', encoding='utf-8') as f:
    for j in range(180, len(doc.paragraphs)):
        f.write(f"[{j}]: style={doc.paragraphs[j].style.name} text='{doc.paragraphs[j].text}'\n")
print("Done")

import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    
    print("=== Paragraph 79 ===")
    print(doc.paragraphs[79].text)
    
    print("\n=== Paragraph 94 ===")
    print(doc.paragraphs[94].text)
    
    print("\n=== Paragraph 157 ===")
    print(doc.paragraphs[157].text)
    
    print("\n=== Paragraph 226 ===")
    print(doc.paragraphs[226].text)
    
    # Search tables for kemahasiswaan
    print("\n=== Searching tables for 'kemahasiswaan' ===")
    for idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "kemahasiswaan" in cell.text.lower():
                    print(f"Table {idx}, Row {r_idx}, Col {c_idx}: '{cell.text}'")

if __name__ == '__main__':
    main()

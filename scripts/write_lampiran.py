import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)

def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.5):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_lampiran_heading(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=6)
    p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def add_lampiran_subheading(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6, line_spacing=1.0)
    p.paragraph_format.left_indent = Inches(0.5)
    
    run = p.add_run(text)
    set_font(run, name="Courier New", size=10)
    
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd {} w:fill="F4F4F4"/>'.format(nsdecls('w')))
    pPr.append(shd)
    
    pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="12" w:space="4" w:color="CCCCCC"/></w:pBdr>'.format(nsdecls('w')))
    pPr.append(pBdr)
    return p

def add_lampiran_image(doc, image_path, width_inches=5.5):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p

def add_image_with_caption(doc, img_path, caption_text, width_inches=2.5):
    p_img = doc.add_paragraph()
    format_paragraph(p_img, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(width_inches))
    
    p_cap = doc.add_paragraph()
    format_paragraph(p_cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6)
    run_cap = p_cap.add_run(caption_text)
    set_font(run_cap, size=10, bold=True)
    return p_img

def main():
    filename = 'Naskah TA Sulthan Syahrul BAB I-IV.docx'
    print(f"Loading {filename}...")
    doc = docx.Document(filename)
    
    # Find paragraph "LAMPIRAN"
    idx_lampiran = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "LAMPIRAN":
            idx_lampiran = idx
            break
            
    if idx_lampiran == -1:
        print("LAMPIRAN heading not found!")
        return
        
    print(f"Found LAMPIRAN heading at index {idx_lampiran}.")
    
    # Delete all paragraphs after idx_lampiran to rebuild from scratch
    p_lampiran = doc.paragraphs[idx_lampiran]
    p_parent = p_lampiran._element.getparent()
    to_delete = doc.paragraphs[idx_lampiran + 1:]
    print(f"Deleting {len(to_delete)} paragraphs after LAMPIRAN heading...")
    for p in to_delete:
        p_parent.remove(p._element)
        
    # Clean up tables
    print(f"Current tables count in document: {len(doc.tables)}")

    # 1. Lampiran A: Manual Penggunaan Aplikasi
    print("Writing Lampiran A (Manual Penggunaan)...")
    add_lampiran_heading(doc, "Lampiran A. Manual Penggunaan Aplikasi")
    
    add_lampiran_subheading(doc, "Lampiran A.1. Panduan Login dan Seleksi Peran")
    add_body(doc,
             "Aplikasi SIMONAS mewajibkan otentikasi pengguna untuk mengamankan data akademik. "
             "Langkah-langkah untuk masuk ke sistem adalah sebagai berikut:")
    add_body(doc, "1. Buka aplikasi SIMONAS di perangkat Android atau melalui peramban web. Masukkan alamat email resmi sekolah yang telah terdaftar (contoh: surono@smpn1jenar.local) dan masukkan kata sandi awal sementara Anda. Setelah itu, tekan tombol 'Login' untuk memproses otentikasi.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_login.png", "Gambar A.1. Tampilan Halaman Login (Placeholder)")
    
    add_body(doc, "2. Apabila akun Anda terdaftar dengan beberapa peran (multi-role), halaman seleksi peran akan muncul secara otomatis. Pilih peran yang ingin Anda gunakan (Kepala Sekolah / Guru / BK / Siswa) untuk dialihkan ke dashboard yang sesuai.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_role_selection.png", "Gambar A.2. Tampilan Halaman Seleksi Peran (Placeholder)")

    add_lampiran_subheading(doc, "Lampiran A.2. Panduan Penggunaan Modul Kepala Sekolah")
    add_body(doc,
             "Modul Kepala Sekolah berfokus pada pengawasan akademik harian. Panduan menu utamanya:")
    
    add_body(doc, "1. Dashboard Utama: Menyajikan ringkasan real-time kehadiran guru hari ini dan status kelengkapan berkas RPP.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_kepsek.png", "Gambar A.3. Tampilan Halaman Dashboard Kepala Sekolah (Placeholder)")
    
    add_body(doc, "2. Monitoring Kelas: Klik menu ini untuk melihat daftar kelas aktif (misal: 7A, 7B). Sistem menampilkan nama guru yang sedang mengajar, mata pelajaran, dan status keterisian jurnal kelas.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_kepsek_monitoring_kelas.png", "Gambar A.4. Tampilan Dashboard Kepala Sekolah dengan Sorotan Menu Monitoring Kelas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_monitoring_kelas.png", "Gambar A.5. Tampilan Halaman Monitoring Kelas oleh Kepala Sekolah (Placeholder)")
    
    add_body(doc, "3. Monitoring Jurnal: Klik untuk melihat riwayat isi jurnal mengajar harian guru secara lengkap beserta materi pelajaran.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_kepsek_monitoring_jurnal.png", "Gambar A.6. Tampilan Dashboard Kepala Sekolah dengan Sorotan Menu Monitoring Jurnal (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_monitoring_jurnal.png", "Gambar A.7. Tampilan Halaman Monitoring Jurnal oleh Kepala Sekolah (Placeholder)")
    
    add_body(doc, "4. Verifikasi Administrasi: Kepala Sekolah dapat melihat berkas perangkat pembelajaran yang diunggah oleh guru. Tekan tombol 'Setujui' atau 'Tolak' untuk memberikan status verifikasi.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_kepsek_verifikasi_administrasi.png", "Gambar A.8. Tampilan Dashboard Kepala Sekolah dengan Sorotan Menu Verifikasi Administrasi (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_verifikasi_administrasi.png", "Gambar A.9. Tampilan Halaman Verifikasi Administrasi oleh Kepala Sekolah (Placeholder)")

    add_lampiran_subheading(doc, "Lampiran A.3. Panduan Penggunaan Modul Guru")
    add_body(doc,
             "Modul Guru digunakan untuk mencatat aktivitas KBM dan mengunggah administrasi pembelajaran. Panduan menu utamanya:")
    
    add_body(doc, "1. Jadwal Mengajar: Menampilkan daftar jadwal mengajar guru berdasarkan hari aktif.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_guru_jadwal_mengajar.png", "Gambar A.10. Tampilan Dashboard Guru dengan Sorotan Menu Jadwal Mengajar (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_jadwal_mengajar.png", "Gambar A.11. Tampilan Halaman Jadwal Mengajar Guru (Placeholder)")
    
    add_body(doc, "2. Input Jurnal and Absensi: Pilih kelas dan mata pelajaran. Ketikkan materi ajar dan catatan. Tandai kehadiran siswa pada daftar absensi (Hadir/Sakit/Izin/Alfa) lalu tekan tombol 'Simpan Jurnal'.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_guru_jurnal_absensi.png", "Gambar A.12. Tampilan Dashboard Guru dengan Sorotan Menu Jurnal & Absensi (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_jurnal_absensi.png", "Gambar A.13. Tampilan Halaman Input Jurnal dan Presensi oleh Guru (Placeholder)")
    
    add_body(doc, "3. Kelola Nilai: Guru dapat menginputkan nilai tugas, UTS, dan UAS siswa secara berkala per kelas.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_guru_kelola_nilai.png", "Gambar A.14. Tampilan Dashboard Guru dengan Sorotan Menu Kelola Nilai (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_kelola_nilai.png", "Gambar A.15. Tampilan Halaman Kelola Nilai Siswa oleh Guru (Placeholder)")
    
    add_body(doc, "4. Kelola Tugas: Menu untuk mempublikasikan tugas baru bagi siswa lengkap dengan batas waktu (deadline).")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_guru_kelola_tugas.png", "Gambar A.16. Tampilan Dashboard Guru dengan Sorotan Menu Kelola Tugas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_kelola_tugas.png", "Gambar A.17. Tampilan Halaman Kelola Tugas Siswa oleh Guru (Placeholder)")
    
    add_body(doc, "5. Berkas Pembelajaran: Unggah file administrasi pembelajaran (format PDF) seperti RPP atau Silabus untuk diverifikasi oleh Kepala Sekolah.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_guru_administrasi.png", "Gambar A.18. Tampilan Dashboard Guru dengan Sorotan Menu Berkas Pembelajaran (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_administrasi.png", "Gambar A.19. Tampilan Halaman Unggah Administrasi Pembelajaran oleh Guru (Placeholder)")

    add_lampiran_subheading(doc, "Lampiran A.4. Panduan Penggunaan Modul Guru BK")
    add_body(doc,
             "Modul Guru BK memfasilitasi pembinaan kesiswaan dan pemetaan kelas. Panduan menu utamanya:")
    
    add_body(doc, "1. Kelola Akun Siswa: BK dapat menambahkan data siswa baru (Nama, NIS, Email) ke dalam sistem.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_bk_kelola_siswa.png", "Gambar A.20. Tampilan Dashboard BK dengan Sorotan Menu Kelola Akun Siswa (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_kelola_siswa.png", "Gambar A.21. Tampilan Halaman Kelola Akun Siswa oleh Guru BK (Placeholder)")
    
    add_body(doc, "2. Pengaturan Kelas: BK memetakan siswa ke dalam kelas-kelas yang aktif pada tahun ajaran berjalan.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_bk_pengaturan_kelas.png", "Gambar A.22. Tampilan Dashboard BK dengan Sorotan Menu Pengaturan Kelas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_pengaturan_kelas.png", "Gambar A.23. Tampilan Halaman Pengaturan Kelas Siswa oleh Guru BK (Placeholder)")
    
    add_body(doc, "3. Kenaikan Kelas: Menu tahunan untuk memproses kenaikan kelas siswa aktif ke tingkat kelas yang lebih tinggi.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_bk_kenaikan_kelas.png", "Gambar A.24. Tampilan Dashboard BK dengan Sorotan Menu Kenaikan Kelas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_kenaikan_kelas.png", "Gambar A.25. Tampilan Halaman Kelola Kenaikan Kelas oleh Guru BK (Placeholder)")
    
    add_body(doc, "4. Catatan Pembinaan: Form untuk mencatat konseling dan pembinaan bagi siswa bermasalah.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_bk_catatan_pembinaan.png", "Gambar A.26. Tampilan Dashboard BK dengan Sorotan Menu Catatan Pembinaan (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_catatan_pembinaan.png", "Gambar A.27. Tampilan Halaman Catatan Pembinaan Siswa oleh Guru BK (Placeholder)")

    add_lampiran_subheading(doc, "Lampiran A.5. Panduan Penggunaan Modul Siswa")
    add_body(doc,
             "Siswa dapat memantau jadwal, tugas, dan rekam akademik pribadi. Panduan menu utamanya:")
    
    add_body(doc, "1. Jadwal Pelajaran: Menampilkan jadwal mata pelajaran mingguan untuk kelas siswa bersangkutan.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_siswa_jadwal_pelajaran.png", "Gambar A.28. Tampilan Dashboard Siswa dengan Sorotan Menu Jadwal Pelajaran (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_jadwal_pelajaran.png", "Gambar A.29. Tampilan Halaman Jadwal Pelajaran Siswa (Placeholder)")
    
    add_body(doc, "2. Nilai Siswa: Menampilkan informasi perolehan nilai tugas, UTS, dan UAS secara transparan.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_siswa_nilai_siswa.png", "Gambar A.30. Tampilan Dashboard Siswa dengan Sorotan Menu Nilai Siswa (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_nilai_siswa.png", "Gambar A.31. Tampilan Halaman Nilai Akademik Siswa (Placeholder)")
    
    add_body(doc, "3. Tugas Kelas: Menampilkan daftar tugas dari guru. Siswa dapat mengunggah file jawaban tugas secara langsung.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_siswa_tugas_kelas.png", "Gambar A.32. Tampilan Dashboard Siswa dengan Sorotan Menu Tugas Kelas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_tugas_kelas.png", "Gambar A.33. Tampilan Halaman Tugas Kelas Siswa (Placeholder)")
    
    add_body(doc, "4. Histori Kelas: Menampilkan riwayat kelas dan wali kelas yang pernah ditempati siswa di masa lalu.")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_dashboard_siswa_histori_kelas.png", "Gambar A.34. Tampilan Dashboard Siswa dengan Sorotan Menu Histori Kelas (Placeholder)")
    add_image_with_caption(doc, "scripts/placeholders/placeholder_histori_kelas.png", "Gambar A.35. Tampilan Halaman Histori Kelas Siswa (Placeholder)")

    # 2. Lampiran B: ERD Modul Pengelolaan Kelas dan Histori Wali Kelas
    print("Writing Lampiran B (ERD)...")
    add_lampiran_heading(doc, "Lampiran B. ERD Modul Pengelolaan Kelas dan Histori Wali Kelas")
    
    add_lampiran_subheading(doc, "Lampiran B.1. ERD Modul Pengelolaan Kelas")
    add_body(doc,
             "Diagram ERD (Entity Relationship Diagram) berikut memetakan hubungan entitas data dalam modul pengelolaan kelas siswa, "
             "yang menghubungkan data siswa dengan pembagian kelas dan tahun akademik aktif di database public.")
    add_lampiran_image(doc, "scripts/extracted_kmm_media/image17.png", width_inches=6.0)
    
    add_lampiran_subheading(doc, "Lampiran B.2. ERD Modul Histori Wali Kelas")
    add_body(doc,
             "Diagram ERD berikut memetakan relasi data untuk pencatatan histori penugasan wali kelas "
             "dari tahun ajaran ke tahun ajaran berikutnya, memastikan rekam data kepegawaian guru tersimpan secara terstruktur.")
    add_lampiran_image(doc, "scripts/extracted_kmm_media/image18.png", width_inches=6.0)

    # 3. Lampiran C: Kebijakan Keamanan Supabase Row Level Security (RLS) dan Trigger
    print("Writing Lampiran C (SQL RLS)...")
    add_lampiran_heading(doc, "Lampiran C. Kebijakan Keamanan Supabase Row Level Security (RLS) dan Trigger")

    tables_rls = [
        {
            "name": "users",
            "title": "Tabel Profil Pengguna (users)",
            "desc": "Berikut adalah skrip SQL yang digunakan untuk mendefinisikan kebijakan keamanan data pada tabel pengguna (public.users) di database Supabase PostgreSQL. Kebijakan ini memastikan bahwa setiap pengguna terautentikasi dapat membaca data pengguna, dan pengguna anonim dapat mencari username sebelum login, serta pembaruan profil hanya bisa dilakukan oleh pemilik profil itu sendiri:",
            "sql": """-- Kebijakan membaca profil oleh pengguna terautentikasi (Select)
CREATE POLICY "Allow select for authenticated users" 
ON public.users
FOR SELECT 
TO authenticated 
USING (true);

-- Kebijakan membaca profil oleh pengguna anonim (Select untuk login)
CREATE POLICY "Allow anonymous select for username lookup" 
ON public.users
FOR SELECT 
TO anon 
USING (true);

-- Kebijakan memperbarui profil sendiri (Update)
CREATE POLICY "Allow update for users own profile" 
ON public.users
FOR UPDATE 
TO authenticated 
USING (auth.uid() = id)
WITH CHECK (auth.uid() = id);"""
        },
        {
            "name": "roles",
            "title": "Tabel Peran (roles)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel roles. Seluruh pengguna terautentikasi dapat membaca data peran, sedangkan penambahan atau perubahan peran hanya dapat dilakukan oleh Kepala Sekolah selaku administrator sistem:",
            "sql": """-- Kebijakan membaca data peran (Select)
CREATE POLICY "Allow select roles for authenticated" 
ON public.roles
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola data peran oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage roles for kepsek" 
ON public.roles
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "user_roles",
            "title": "Tabel Relasi Peran Pengguna (user_roles)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel user_roles. Relasi peran pengguna dapat dibaca oleh seluruh pengguna terautentikasi untuk otorisasi akses menu di aplikasi mobile, sedangkan perubahan relasi peran hanya diizinkan untuk Kepala Sekolah:",
            "sql": """-- Kebijakan membaca data peran pengguna (Select)
CREATE POLICY "Allow select user_roles for authenticated" 
ON public.user_roles
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola peran pengguna oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage user_roles for kepsek" 
ON public.user_roles
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "guru",
            "title": "Tabel Profil Guru (guru)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel profil guru (public.guru). Seluruh pengguna terautentikasi dapat melihat profil guru, guru dapat memperbarui datanya sendiri, serta Kepala Sekolah dan staf Kemahasiswaan memiliki akses penuh untuk mengelola data kepegawaian guru:",
            "sql": """-- Kebijakan membaca data guru (Select)
CREATE POLICY "Allow select guru for authenticated" 
ON public.guru
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan memperbarui data mandiri oleh Guru (Update)
CREATE POLICY "Allow update own data for guru" 
ON public.guru
FOR UPDATE 
TO authenticated 
USING (auth.uid() = user_id)
WITH CHECK (auth.uid() = user_id);

-- Kebijakan mengelola data guru oleh Kepsek dan Kemahasiswaan (All)
CREATE POLICY "Allow manage guru for kepsek and kemahasiswaan" 
ON public.guru
FOR ALL 
TO authenticated 
USING (
  user_has_role(auth.uid(), 'kepsek') OR 
  user_has_role(auth.uid(), 'kemahasiswaan')
);"""
        },
        {
            "name": "siswa",
            "title": "Tabel Profil Siswa (siswa)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel profil siswa (public.siswa). Seluruh pengguna terautentikasi dapat membaca data siswa untuk kebutuhan interaksi aplikasi, sedangkan pendaftaran dan perubahan data siswa hanya dikelola oleh Kepala Sekolah dan staf Kemahasiswaan:",
            "sql": """-- Kebijakan membaca data siswa (Select)
CREATE POLICY "Allow select siswa for authenticated" 
ON public.siswa
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola data siswa oleh Kepsek dan Kemahasiswaan (All)
CREATE POLICY "Allow manage siswa for kepsek and kemahasiswaan" 
ON public.siswa
FOR ALL 
TO authenticated 
USING (
  user_has_role(auth.uid(), 'kepsek') OR 
  user_has_role(auth.uid(), 'kemahasiswaan')
);"""
        },
        {
            "name": "tahun_ajaran",
            "title": "Tabel Tahun Ajaran (tahun_ajaran)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel tahun_ajaran. Semua pengguna terautentikasi dapat melihat daftar tahun ajaran berjalan, sedangkan pengelolaan tahun ajaran baru atau pengaktifan semester dikonfigurasi oleh Kepala Sekolah:",
            "sql": """-- Kebijakan membaca data tahun ajaran (Select)
CREATE POLICY "Allow select tahun_ajaran for authenticated" 
ON public.tahun_ajaran
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola tahun ajaran oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage tahun_ajaran for kepsek" 
ON public.tahun_ajaran
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "semester",
            "title": "Tabel Semester (semester)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel semester. Akses baca diberikan kepada seluruh pengguna terautentikasi, sedangkan perubahan konfigurasi semester aktif dibatasi khusus untuk Kepala Sekolah selaku pimpinan akademik:",
            "sql": """-- Kebijakan membaca semester (Select)
CREATE POLICY "Allow select semester for authenticated" 
ON public.semester
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola data semester oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage semester for kepsek" 
ON public.semester
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "kelas",
            "title": "Tabel Kelas (kelas)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel kelas. Semua pengguna terautentikasi dapat melihat daftar kelas, sedangkan pembuatan, perubahan nama kelas, dan penentuan tingkat kelas dikelola oleh Kepala Sekolah dan staf Kemahasiswaan:",
            "sql": """-- Kebijakan membaca kelas (Select)
CREATE POLICY "Allow select kelas for authenticated" 
ON public.kelas
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola kelas oleh Kepsek dan Kemahasiswaan (All)
CREATE POLICY "Allow manage kelas for kepsek and kemahasiswaan" 
ON public.kelas
FOR ALL 
TO authenticated 
USING (
  user_has_role(auth.uid(), 'kepsek') OR 
  user_has_role(auth.uid(), 'kemahasiswaan')
);"""
        },
        {
            "name": "mapel",
            "title": "Tabel Mata Pelajaran (mapel)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel mapel (mata pelajaran). Seluruh pengguna terautentikasi dapat melihat daftar mata pelajaran yang aktif di sekolah, sedangkan penambahan mata pelajaran baru dikelola oleh Kepala Sekolah:",
            "sql": """-- Kebijakan membaca mata pelajaran (Select)
CREATE POLICY "Allow select mapel for authenticated" 
ON public.mapel
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola mapel oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage mapel for kepsek" 
ON public.mapel
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "jadwal",
            "title": "Tabel Jadwal Pelajaran (jadwal)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel jadwal pelajaran. Guru pengajar dapat melihat jadwal mengajar mereka sendiri, siswa melihat jadwal pelajaran di kelasnya yang aktif, sedangkan Kepala Sekolah dan staf Kemahasiswaan dapat memantau seluruh jadwal dan mengelolanya secara penuh:",
            "sql": """-- Kebijakan membaca jadwal pelajaran (Select)
CREATE POLICY "Allow select jadwal based on role" 
ON public.jadwal
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  user_has_role(auth.uid(), 'kemahasiswaan') OR
  EXISTS (
    SELECT 1 FROM public.kelas_siswa ks
    WHERE ks.siswa_uid = (SELECT id FROM public.siswa WHERE user_id = auth.uid())
    AND ks.kelas_id = jadwal.kelas_id
    AND ks.status_aktif = true
  )
)
AND is_deleted = false;

-- Kebijakan mengelola jadwal pelajaran oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage jadwal for kepsek" 
ON public.jadwal
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        },
        {
            "name": "jurnal",
            "title": "Tabel Jurnal Mengajar (jurnal)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel jurnal mengajar harian. Guru yang mengajar berhak mencatat dan memperbarui jurnal mengajar kelasnya, siswa berhak melihat riwayat jurnal pelajaran untuk kelas mereka, sedangkan Kepala Sekolah dan Kemahasiswaan dapat mengakses seluruh jurnal untuk kebutuhan pengawasan:",
            "sql": """-- Kebijakan membaca jurnal (Select)
CREATE POLICY "Allow select jurnal based on role" 
ON public.jurnal
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  user_has_role(auth.uid(), 'kemahasiswaan') OR
  EXISTS (
    SELECT 1 FROM public.kelas_siswa ks
    WHERE ks.siswa_uid = (SELECT id FROM public.siswa WHERE user_id = auth.uid())
    AND ks.kelas_id = jurnal.kelas_id
    AND ks.status_aktif = true
  )
)
AND is_deleted = false;

-- Kebijakan mengelola jurnal (All) oleh Guru pengajar yang bersangkutan
CREATE POLICY "Allow manage own jurnal for guru" 
ON public.jurnal
FOR ALL 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid)
);"""
        },
        {
            "name": "absensi_jurnal",
            "title": "Tabel Presensi Siswa (absensi_jurnal)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel absensi_jurnal (kehadiran siswa). Guru pengajar dapat mengisi dan memperbarui data absensi siswa, siswa dapat melihat status kehadiran mereka sendiri, sedangkan Kepala Sekolah dan Kemahasiswaan memiliki hak akses pemantauan secara menyeluruh:",
            "sql": """-- Kebijakan membaca absensi (Select)
CREATE POLICY "Allow select absensi_jurnal based on role" 
ON public.absensi_jurnal
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.siswa WHERE id = siswa_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  EXISTS (
    SELECT 1 FROM public.jurnal j
    WHERE j.id = absensi_jurnal.jurnal_id
    AND auth.uid() = (SELECT user_id FROM public.guru WHERE id = j.guru_uid)
  )
)
AND is_deleted = false;

-- Kebijakan mengelola absensi (All) oleh Guru pengajar kelas
CREATE POLICY "Allow manage absensi for teaching guru" 
ON public.absensi_jurnal
FOR ALL 
TO authenticated 
USING (
  EXISTS (
    SELECT 1 FROM public.jurnal j
    WHERE j.id = absensi_jurnal.jurnal_id
    AND auth.uid() = (SELECT user_id FROM public.guru WHERE id = j.guru_uid)
  )
);"""
        },
        {
            "name": "kelas_siswa",
            "title": "Tabel Pembagian Kelas Siswa (kelas_siswa)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel pembagian kelas siswa (kelas_siswa). Seluruh pengguna terautentikasi dapat melihat penempatan kelas siswa, sedangkan pengelolaan penempatan kelas dan mutasi siswa dikelola oleh staf Kemahasiswaan dan Kepala Sekolah:",
            "sql": """-- Kebijakan membaca pembagian kelas (Select)
CREATE POLICY "Allow select kelas_siswa for authenticated" 
ON public.kelas_siswa
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola penempatan kelas siswa oleh Kepsek dan Kemahasiswaan (All)
CREATE POLICY "Allow manage kelas_siswa for kepsek and kemahasiswaan" 
ON public.kelas_siswa
FOR ALL 
TO authenticated 
USING (
  user_has_role(auth.uid(), 'kepsek') OR 
  user_has_role(auth.uid(), 'kemahasiswaan')
);"""
        },
        {
            "name": "nilai_ujian",
            "title": "Tabel Nilai Ujian Siswa (nilai_ujian)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel nilai_ujian (UTS/UAS). Guru pengajar dapat melihat, menambah, dan memperbarui nilai siswa, siswa yang bersangkutan hanya dapat melihat nilai mereka sendiri, sedangkan Kepala Sekolah dapat memantau seluruh nilai ujian untuk pelaporan:",
            "sql": """-- Kebijakan membaca nilai ujian (Select)
CREATE POLICY "Allow select nilai_ujian based on role" 
ON public.nilai_ujian
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.siswa WHERE id = siswa_uid) OR
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid) OR
  user_has_role(auth.uid(), 'kepsek')
)
AND is_deleted = false;

-- Kebijakan mengelola nilai ujian (All) oleh Guru pengajar mata pelajaran
CREATE POLICY "Allow manage nilai_ujian for teaching guru" 
ON public.nilai_ujian
FOR ALL 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid)
);"""
        },
        {
            "name": "tugas",
            "title": "Tabel Tugas Siswa (tugas)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel tugas. Guru pengajar dapat mengelola (membuat, mengubah, menghapus) tugas untuk kelas mereka. Siswa berhak melihat daftar tugas untuk kelasnya, dan Kepala Sekolah memantau tugas yang diberikan:",
            "sql": """-- Kebijakan membaca tugas (Select)
CREATE POLICY "Allow select tugas based on role" 
ON public.tugas
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  EXISTS (
    SELECT 1 FROM public.kelas_siswa ks
    WHERE ks.siswa_uid = (SELECT id FROM public.siswa WHERE user_id = auth.uid())
    AND ks.kelas_id = tugas.kelas_id
    AND ks.status_aktif = true
  )
)
AND is_deleted = false;

-- Kebijakan mengelola tugas (All) oleh Guru pembuat tugas
CREATE POLICY "Allow manage tugas for creator guru" 
ON public.tugas
FOR ALL 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid)
);"""
        },
        {
            "name": "nilai_tugas",
            "title": "Tabel Nilai Tugas Siswa (nilai_tugas)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel nilai_tugas. Guru pengajar dapat mengelola nilai tugas siswa. Siswa berhak melihat nilai tugas miliknya sendiri, sedangkan Kepala Sekolah dapat memantau seluruh nilai tugas siswa:",
            "sql": """-- Kebijakan membaca nilai tugas (Select)
CREATE POLICY "Allow select nilai_tugas based on role" 
ON public.nilai_tugas
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.siswa WHERE id = siswa_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  EXISTS (
    SELECT 1 FROM public.tugas t
    WHERE t.id = nilai_tugas.tugas_id
    AND auth.uid() = (SELECT user_id FROM public.guru WHERE id = t.guru_uid)
  )
)
AND is_deleted = false;

-- Kebijakan mengelola nilai tugas (All) oleh Guru pengajar kelas
CREATE POLICY "Allow manage nilai_tugas for teaching guru" 
ON public.nilai_tugas
FOR ALL 
TO authenticated 
USING (
  EXISTS (
    SELECT 1 FROM public.tugas t
    WHERE t.id = nilai_tugas.tugas_id
    AND auth.uid() = (SELECT user_id FROM public.guru WHERE id = t.guru_uid)
  )
);"""
        },
        {
            "name": "administrasi_pembelajaran",
            "title": "Tabel Administrasi Mengajar Guru (administrasi_pembelajaran)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel berkas administrasi_pembelajaran guru. Guru yang bersangkutan berhak mengelola dokumen (RPP, Silabus, Prota, Promes) miliknya sendiri. Kepala Sekolah dan Kemahasiswaan memiliki hak penuh untuk membaca seluruh dokumen guna keperluan verifikasi dan monitoring:",
            "sql": """-- Kebijakan membaca berkas administrasi (Select)
CREATE POLICY "Allow select administrasi based on role" 
ON public.administrasi_pembelajaran
FOR SELECT 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid) OR
  user_has_role(auth.uid(), 'kepsek') OR
  user_has_role(auth.uid(), 'kemahasiswaan')
)
AND is_deleted = false;

-- Kebijakan mengelola berkas administrasi oleh Guru pengunggah (All)
CREATE POLICY "Allow manage own administrasi for guru" 
ON public.administrasi_pembelajaran
FOR ALL 
TO authenticated 
USING (
  auth.uid() = (SELECT user_id FROM public.guru WHERE id = guru_uid)
);"""
        },
        {
            "name": "histori_wali_kelas",
            "title": "Tabel Riwayat Wali Kelas (histori_wali_kelas)",
            "desc": "Berikut adalah skrip SQL untuk kebijakan keamanan tabel histori_wali_kelas. Seluruh pengguna terautentikasi dapat membaca data riwayat penugasan wali kelas sekolah, sementara pencatatan data riwayat penugasan baru dikelola oleh Kepala Sekolah:",
            "sql": """-- Kebijakan membaca riwayat wali kelas (Select)
CREATE POLICY "Allow select histori_wali_kelas for authenticated" 
ON public.histori_wali_kelas
FOR SELECT 
TO authenticated 
USING (is_deleted = false);

-- Kebijakan mengelola riwayat wali kelas oleh Kepala Sekolah (All)
CREATE POLICY "Allow manage histori_wali_kelas for kepsek" 
ON public.histori_wali_kelas
FOR ALL 
TO authenticated 
USING (user_has_role(auth.uid(), 'kepsek'));"""
        }
    ]

    for idx, tbl in enumerate(tables_rls):
        num = idx + 1
        add_lampiran_subheading(doc, f"Lampiran C.{num}. SQL Kebijakan Row Level Security (RLS) {tbl['title']}")
        add_body(doc, tbl['desc'])
        add_code_block(doc, tbl['sql'])

    # Add the trigger as Lampiran C.19
    add_lampiran_subheading(doc, "Lampiran C.19. SQL Trigger Sinkronisasi Akun Auth dengan Tabel Profil")
    add_body(doc,
             "Trigger database berikut berfungsi untuk menyinkronkan data pengguna secara otomatis saat akun baru didaftarkan "
             "melalui modul pendaftaran Supabase Auth (auth.users) ke tabel profil publik (public.users) untuk kebutuhan operasional aplikasi:")
    
    sql_trigger = (
        "-- Fungsi sinkronisasi akun baru\n"
        "CREATE OR REPLACE FUNCTION public.handle_new_user()\n"
        "RETURNS trigger AS $$\n"
        "BEGIN\n"
        "  INSERT INTO public.users (id, email, nama, role)\n"
        "  VALUES (\n"
        "    new.id,\n"
        "    new.email,\n"
        "    coalesce(new.raw_user_meta_data->>'nama', 'Pengguna Baru'),\n"
        "    coalesce(new.raw_user_meta_data->>'role', 'siswa')\n"
        "  );\n"
        "  RETURN new;\n"
        "END;\n"
        "$$ LANGUAGE plpgsql SECURITY DEFINER;\n\n"
        "-- Trigger pemicu sinkronisasi setelah registrasi auth\n"
        "CREATE OR REPLACE TRIGGER on_auth_user_created\n"
        "  AFTER INSERT ON auth.users\n"
        "  FOR EACH ROW EXECUTE FUNCTION public.handle_new_user();"
    )
    add_code_block(doc, sql_trigger)

    # Save document
    print(f"Saving changes to {filename}...")
    try:
        doc.save(filename)
        print("Document saved successfully!")
    except PermissionError:
        fallback = filename.replace(".docx", "_updated.docx")
        print(f"Permission denied. Saving to fallback file: {fallback}")
        doc.save(fallback)
        print(f"Document saved successfully as '{fallback}'!")

if __name__ == '__main__':
    main()

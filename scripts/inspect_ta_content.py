import docx

doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
with open('scripts/inspect_ta_content.txt', 'w', encoding='utf-8') as f:
    for idx, p in enumerate(doc.paragraphs):
        f.write(f"[{idx}] (Style={p.style.name}): {p.text}\n")
print(f"Done. Wrote {len(doc.paragraphs)} paragraphs.")

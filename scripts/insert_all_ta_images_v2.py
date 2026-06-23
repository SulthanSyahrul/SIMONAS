import os
import glob
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc_path = "Naskah_TA_SinarBelajar_SMPN1Jenar_v3_new_sitasi.docx"
    print(f"Opening {doc_path}...")
    doc = docx.Document(doc_path)
    
    brain_dir = r"C:\Users\LENOVO\.gemini\antigravity-ide\brain\b076bab6-1987-420e-bb6c-7db39a4397de"
    
    def get_latest_file(pattern):
        files = glob.glob(os.path.join(brain_dir, pattern))
        if not files:
            print(f"Warning: No files found matching pattern: {pattern}")
            return None
        # Sort by modification time descending
        files.sort(key=os.path.getmtime, reverse=True)
        return files[0]
        
    mapping = {
        # UML Diagrams (Refined / Updated)
        "Gambar 1.1. Kondisi proses akademik sebelum digitalisasi": ("scripts/extracted_images/diagram_manual_process.png", 5.8),
        "Gambar 2.1. Gambaran umum produk SIMONAS": ("scripts/extracted_images/diagram_sinarbelajar_overview.png", 5.8),
        "Gambar 3.3. Activity diagram login dan pemilihan role": ("scripts/extracted_images/diagram_activity_login.png", 5.5),
        "Gambar 3.4. Activity diagram input jurnal dan absensi": ("scripts/extracted_images/diagram_activity_jurnal.png", 5.5),
        
        # Screenshots (Captured from Samsung Galaxy S8 simulation)
        "Gambar 4.1. Implementasi halaman login": (get_latest_file("gambar_4_1_*.png"), 3.2),
        "Gambar 4.2. Implementasi halaman pemilihan role": (get_latest_file("gambar_4_2_*.png"), 3.2),
        "Gambar 4.3. Implementasi dashboard guru": (get_latest_file("gambar_4_3_*.png"), 3.2),
        "Gambar 4.4. Implementasi jurnal dan absensi": (get_latest_file("gambar_4_4_*.png"), 3.2),
        "Gambar 4.5. Implementasi dashboard kepala sekolah": (get_latest_file("gambar_4_5_*.png"), 3.2),
        "Gambar 4.6. Implementasi monitoring kelas": (get_latest_file("gambar_4_6_*.png"), 3.2),
        "Gambar 4.7. Implementasi dashboard BK": (get_latest_file("gambar_4_7_*.png"), 3.2),
        "Gambar 4.8. Implementasi dashboard siswa": (get_latest_file("gambar_4_8_*.png"), 3.2),
        "Gambar 4.9. Implementasi administrasi pembelajaran": (get_latest_file("gambar_4_9_*.png"), 3.2)
    }
    
    # Find all paragraph elements matching mapping captions
    paragraphs_to_process = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in mapping:
            paragraphs_to_process.append((p._element, text))
            
    for p_elem, text in paragraphs_to_process:
        img_path, width_inch = mapping[text]
        if not img_path:
            print(f"Warning: No valid image path for caption '{text}', skipping.")
            continue
        if not os.path.exists(img_path):
            print(f"Warning: Image file not found: {img_path}, skipping.")
            continue
            
        print(f"\nProcessing caption: '{text}'")
        
        # Locate paragraph index in the current list of paragraphs
        idx = -1
        for i, curr_p in enumerate(doc.paragraphs):
            if curr_p._element == p_elem:
                idx = i
                break
                
        if idx == -1:
            print(f"Error: Could not locate paragraph object in document for '{text}'")
            continue
            
        p = doc.paragraphs[idx]
            
        # Check preceding 3 paragraphs for existing drawings to delete (only for diagrams that need replacement)
        should_delete_old_drawing = text in [
            "Gambar 1.1. Kondisi proses akademik sebelum digitalisasi",
            "Gambar 2.1. Gambaran umum produk SIMONAS",
            "Gambar 3.3. Activity diagram login dan pemilihan role",
            "Gambar 3.4. Activity diagram input jurnal dan absensi"
        ]
        
        deleted_count = 0
        if should_delete_old_drawing:
            # Check idx-1, idx-2, idx-3
            for offset in [1, 2, 3]:
                check_idx = idx - offset
                if 0 <= check_idx < len(doc.paragraphs):
                    p_check = doc.paragraphs[check_idx]
                    if 'w:drawing' in p_check._p.xml:
                        print(f"  Deleting existing drawing paragraph at index {check_idx}")
                        p_check._element.getparent().remove(p_check._element)
                        deleted_count += 1
                    
        # Find index again since we deleted paragraphs
        for i, curr_p in enumerate(doc.paragraphs):
            if curr_p._element == p_elem:
                idx = i
                break
                
        print(f"  Inserting new image paragraph before the caption (current index {idx})")
        # Insert image paragraph
        new_p = p.insert_paragraph_before()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the image
        run = new_p.add_run()
        run.add_picture(img_path, width=Inches(width_inch))
        
        # Set format options to look neat
        new_p.paragraph_format.space_before = docx.shared.Pt(12)
        new_p.paragraph_format.space_after = docx.shared.Pt(6)
        
    try:
        doc.save(doc_path)
        print(f"\nSuccessfully saved updated document to {doc_path}!")
    except Exception as e:
        print(f"\nError saving to {doc_path}: {e}")
        backup_path = "Naskah_TA_SinarBelajar_SMPN1Jenar_v3_new_sitasi_updated.docx"
        try:
            doc.save(backup_path)
            print(f"Saved to backup file instead: {backup_path}")
        except Exception as ex:
            print(f"Failed to save to backup as well: {ex}")

if __name__ == "__main__":
    main()

import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    print("Total tables in document:", len(doc.tables))
    
    # We want to verify index 11 to index 37
    # 11: Tabel 5.1
    # 12: Tabel 5.2
    # 13 to 36: Tabel 5.3 to 5.26
    # 37: Tabel 5.27
    
    print("\nVerifying select tables:")
    for idx in range(11, len(doc.tables)):
        table = doc.tables[idx]
        print(f"\n=========================================")
        print(f"Table Index {idx}: Rows={len(table.rows)}, Cols={len(table.columns)}")
        print(f"=========================================")
        # Print first row (header or ID) and a couple of key details
        for r_idx in range(min(5, len(table.rows))):
            row = table.rows[r_idx]
            cell_texts = [f"C{c_idx}: '{cell.text.strip().replace('\n', ' [NL] ')[:60]}'" for c_idx, cell in enumerate(row.cells)]
            print(f"Row {r_idx:02d}: {' | '.join(cell_texts)}")
        if len(table.rows) > 5:
            print(f"... and {len(table.rows) - 5} more rows")

if __name__ == "__main__":
    main()

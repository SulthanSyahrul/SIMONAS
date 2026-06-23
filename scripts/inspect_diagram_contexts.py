import docx

def print_p(doc, start, end):
    print(f"=== Paragraphs {start} to {end} ===")
    for i in range(start, min(len(doc.paragraphs), end + 1)):
        print(f"P {i:03d}: '{doc.paragraphs[i].text}'")

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    print_p(doc, 106, 112)
    print_p(doc, 205, 211)

if __name__ == '__main__':
    main()

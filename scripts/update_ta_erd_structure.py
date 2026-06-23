import docx
import os

def update_document():
    doc_path = "Naskah_TA_SinarBelajar_SMPN1Jenar_v3_cited.docx"
    print(f"Opening {doc_path}...")
    doc = docx.Document(doc_path)
    
    # 1. Update general ERD paragraph
    found_erd_p = False
    for p in doc.paragraphs:
        if "Relasi antarentitas data akademik sekolah digambarkan dalam ERD pada Gambar 3.5" in p.text:
            print("Updating general ERD paragraph reference...")
            p.text = "Relasi antarentitas data akademik sekolah digambarkan dalam ERD pada Gambar 3.5, Gambar 3.6, Gambar 3.7, dan Gambar 3.8. ERD ini memuat relasi primary key dan foreign key dari tabel-tabel di atas."
            found_erd_p = True
            break
    if not found_erd_p:
        print("Warning: General ERD paragraph not found!")

    # 2. Split Gambar 3.5 caption and insert descriptive paragraphs & captions for 3.6, 3.7, 3.8
    # Also insert the transitional paragraph pointing to the appendix ERDs (Lampiran E) at the end.
    found_caption_35 = False
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Gambar 3.5. Entity relationship diagram":
            print(f"Found Gambar 3.5 caption at paragraph {idx}. Splitting...")
            # Modify 3.5 caption
            p.text = "Gambar 3.5. ERD Modul Pengguna"
            
            # The next paragraph is where we will insert our new content before
            p_next = doc.paragraphs[idx + 1]
            
            # Insert spacing and details for Gambar 3.6
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Modul akademik yang berkaitan dengan jadwal pembelajaran, jurnal mengajar, dan absensi dirancang agar proses monitoring akademik dapat dilakukan secara terintegrasi. Relasi antar entitas pada modul tersebut dapat dilihat pada Gambar 3.6.")
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Gambar 3.6. ERD Modul Jadwal dan Jurnal")
            
            # Insert spacing and details for Gambar 3.7
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Modul penilaian dirancang untuk mengelola nilai ujian, tugas, serta ulangan siswa secara terstruktur. Relasi basis data pada modul penilaian dapat dilihat pada Gambar 3.7.")
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Gambar 3.7. ERD Modul Penilaian")
            
            # Insert spacing and details for Gambar 3.8
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Modul administrasi pembelajaran dirancang untuk mengelola dokumen perangkat pembelajaran guru, seperti silabus, Program Tahunan (Prota), Program Semester (Promes), dan Rencana Pelaksanaan Pembelajaran (RPP). Pada modul ini, setiap dokumen dihubungkan dengan data guru, kelas, mata pelajaran, dan tahun ajaran agar penyimpanan administrasi lebih terstruktur. Relasi basis data pada modul administrasi pembelajaran dapat dilihat pada Gambar 3.8.")
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Gambar 3.8. ERD Modul Administrasi Pembelajaran")
            
            # Insert transitional paragraph pointing to Lampiran E
            p_next.insert_paragraph_before("")
            p_next.insert_paragraph_before("Perancangan basis data pada bagian utama hanya menampilkan entitas inti yang mendukung proses utama sistem. Detail relasi data pada modul pengelolaan kelas siswa serta histori wali kelas disajikan pada bagian lampiran dan dapat dilihat pada Lampiran E.")
            
            found_caption_35 = True
            break
            
    if not found_caption_35:
        print("Warning: Caption 'Gambar 3.5. Entity relationship diagram' not found!")

    # 3. Renumber subsequent figures and update their references
    # We find and modify captions and paragraphs
    for p in doc.paragraphs:
        # Renumber captions
        if p.text.strip() == "Gambar 3.6. Rancangan antarmuka login":
            print("Renumbering caption 3.6 -> 3.9")
            p.text = "Gambar 3.9. Rancangan antarmuka login"
        elif p.text.strip() == "Gambar 3.7. Rancangan dashboard multi-role":
            print("Renumbering caption 3.7 -> 3.10")
            p.text = "Gambar 3.10. Rancangan dashboard multi-role"
            
        # Renumber text references
        if "Rancangan antarmuka visual untuk halaman autentikasi masuk ditunjukkan pada Gambar 3.6." in p.text:
            print("Updating text reference to Gambar 3.6 -> 3.9")
            p.text = p.text.replace("Gambar 3.6.", "Gambar 3.9.")
        if "Desain visual halaman utama (dashboard) untuk masing-masing peran pengguna ditunjukkan pada Gambar 3.7." in p.text:
            print("Updating text reference to Gambar 3.7 -> 3.10")
            p.text = p.text.replace("Gambar 3.7.", "Gambar 3.10.")

    # 4. Update the Table of Figures (Daftar Gambar) list
    found_fig_list = False
    for p in doc.paragraphs:
        if "Gambar 3.5 Entity relationship diagram" in p.text:
            print("Updating Table of Figures list...")
            lines = p.text.split("\n")
            new_lines = []
            for line in lines:
                if "Gambar 3.5 Entity relationship diagram" in line:
                    new_lines.append("Gambar 3.5 ERD Modul Pengguna\t16")
                    new_lines.append("Gambar 3.6 ERD Modul Jadwal dan Jurnal\t16")
                    new_lines.append("Gambar 3.7 ERD Modul Penilaian\t16")
                    new_lines.append("Gambar 3.8 ERD Modul Administrasi Pembelajaran\t16")
                elif "Gambar 3.6 Rancangan antarmuka login" in line:
                    new_lines.append("Gambar 3.9 Rancangan antarmuka login\t17")
                elif "Gambar 3.7 Rancangan dashboard multi-role" in line:
                    new_lines.append("Gambar 3.10 Rancangan dashboard multi-role\t17")
                else:
                    new_lines.append(line)
            p.text = "\n".join(new_lines)
            found_fig_list = True
            break
    if not found_fig_list:
        print("Warning: Table of Figures list not found!")

    # 5. Add Lampiran E to the list of appendices at the end of the document
    found_appendix_list = False
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Lampiran D. Skema database SQL lengkap dan kebijakan Row-Level Security (RLS) Supabase.":
            print(f"Found Lampiran D at paragraph {idx}. Appending Lampiran E to list...")
            # Insert Lampiran E in the list
            if idx + 1 < len(doc.paragraphs):
                p_next = doc.paragraphs[idx + 1]
                p_next.insert_paragraph_before("Lampiran E. ERD Modul Pengelolaan Kelas dan Histori Wali Kelas.")
            else:
                doc.add_paragraph("Lampiran E. ERD Modul Pengelolaan Kelas dan Histori Wali Kelas.")
            found_appendix_list = True
            break
    if not found_appendix_list:
        print("Warning: Appendix list item for Lampiran D not found!")

    # 6. Append actual pages/headings and descriptions for Lampiran E at the end of the document
    print("Appending Lampiran E sections to the end of the document...")
    doc.add_paragraph("")
    doc.add_paragraph("Lampiran E. ERD Modul Pengelolaan Kelas dan Histori Wali Kelas")
    doc.add_paragraph("Detail relasi basis data yang digunakan untuk penempatan kelas siswa, riwayat kelas, serta pembagian wali kelas dapat dilihat pada Lampiran E.1.")
    doc.add_paragraph("")
    doc.add_paragraph("Lampiran E.1. ERD Modul Pengelolaan Kelas")
    doc.add_paragraph("Sedangkan riwayat penugasan wali kelas dari tahun ke tahun beserta histori kelas yang diawasi dapat dilihat pada Lampiran E.2.")
    doc.add_paragraph("")
    doc.add_paragraph("Lampiran E.2. ERD Modul Histori Wali Kelas")

    doc.save(doc_path)
    print("Restructuring TA document completed successfully!")

if __name__ == "__main__":
    update_document()

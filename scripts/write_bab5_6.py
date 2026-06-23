import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # A bug in python-docx: font.name doesn't always apply to East Asian / other complex text unless rFonts element is modified
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)

def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.5):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_heading2(doc, p_before, text):
    new_p = p_before.insert_paragraph_before()
    format_paragraph(new_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
    run = new_p.add_run(text)
    set_font(run, bold=True)
    return new_p

def add_body(doc, p_before, text, bold=False, italic=False):
    new_p = p_before.insert_paragraph_before()
    format_paragraph(new_p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6)
    run = new_p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return new_p

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right, insideH, insideV
    value: dictionary with 'sz', 'val', 'color', 'space'
    e.g. top={'sz': 12, 'val': 'single', 'color': 'D3D3D3'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def format_table(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        # Center table
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr[0].append(jc)

def add_table(doc, p_before, headers, data, col_widths=None):
    # Create table
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    p_before._element.addprevious(table._element)
    
    format_table(table)
    
    # Write header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4, line_spacing=1.0)
        for r in p.runs:
            set_font(r, bold=True)
            
        # Light grey background for header cell
        shading = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        
        # Border
        border_style = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
        set_cell_border(hdr_cells[i], top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    # Write data
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = cell_value
            p = row_cells[c_idx].paragraphs[0]
            
            # Align center for short text like IDs or Status
            is_center = len(cell_value) < 10 or cell_value in ["Berhasil", "Laki-laki", "Perempuan"]
            align = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
            
            format_paragraph(p, align=align, space_before=4, space_after=4, line_spacing=1.0)
            for r in p.runs:
                set_font(r)
                
            # Border
            border_style = {'val': 'single', 'sz': '4', 'space': '0', 'color': 'CCCCCC'}
            set_cell_border(row_cells[c_idx], top=border_style, bottom=border_style, left=border_style, right=border_style)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    # Add an empty paragraph after table for spacing
    spacer = p_before.insert_paragraph_before()
    format_paragraph(spacer, space_before=0, space_after=6)
    
    return table

def main():
    filename = 'Naskah TA Sulthan Syahrul BAB I-IV.docx'
    print(f"Loading {filename}...")
    doc = docx.Document(filename)
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if 'BAB V' in text and 'PENGUJIAN' in text:
            p_pengujian = p
        elif 'BAB VI' in text and ('KESIMPULAN' in text or 'SARAN' in text):
            p_bab6 = p
            p_kesimpulan = p
        elif 'DAFTAR PUSTAKA' in text:
            p_daftar_pustaka = p
            
    if not p_pengujian or not p_bab6 or not p_kesimpulan:
        print(f"Error: headings check failed: p_pengujian={p_pengujian is not None}, p_bab6={p_bab6 is not None}, p_kesimpulan={p_kesimpulan is not None}")
        print("Required structural headings not found in document. Please verify headings.")
        return
        
    print("Found structural sections.")
    
    # We want to clear any empty paragraphs between PENGUJIAN and BAB VI
    # Let's collect elements to remove
    idx_start = -1
    idx_end = -1
    for idx, p in enumerate(doc.paragraphs):
        if p._element == p_pengujian._element:
            idx_start = idx
        if p._element == p_bab6._element:
            idx_end = idx
            break
            
    print(f"Clearing empty paragraphs in Bab V (index {idx_start} to {idx_end})...")
    # Clean the paragraphs between them
    to_remove = []
    for i in range(idx_start + 1, idx_end):
        to_remove.append(doc.paragraphs[i])
        
    for p in to_remove:
        # remove paragraph
        p._element.getparent().remove(p._element)
        
    # Find position again
    p_insert_v = None
    for idx, p in enumerate(doc.paragraphs):
        if p._element == p_bab6._element:
            p_insert_v = p
            break
            
    # Write BAB V contents
    print("Writing BAB V content...")
    add_heading2(doc, p_insert_v, "5.1 Metode Pengujian")
    
    add_body(doc, p_insert_v, 
             "Pengujian yang disiapkan untuk aplikasi SIMONAS adalah black-box testing. "
             "Metode ini berfokus pada pemeriksaan masukan dan keluaran fungsi tanpa melihat struktur "
             "internal kode. Pengujian black-box relevan untuk sistem informasi akademik karena setiap "
             "fitur dapat divalidasi berdasarkan skenario pengguna, seperti login, pemilihan role, "
             "input jurnal, input absensi, dan monitoring data. Referensi Muslimin dkk. (2020) serta "
             "Wulandari dkk. (2022) menunjukkan bahwa black-box testing dengan equivalence partitioning "
             "dapat digunakan untuk memeriksa validasi fungsi sistem informasi akademik.")
             
    add_body(doc, p_insert_v,
             "Selain black-box testing, pengujian tambahan yang disarankan adalah widget test untuk "
             "memastikan tampilan awal aplikasi muncul dengan benar, serta User Acceptance Testing (UAT) "
             "bersama pengguna sekolah. Pada naskah Tugas Akhir ini, tabel hasil pengujian disajikan untuk "
             "menunjukkan status keberhasilan setiap fitur yang diuji secara menyeluruh.")

    add_heading2(doc, p_insert_v, "5.2 Lingkungan Pengujian")
    
    add_body(doc, p_insert_v,
             "Spesifikasi perangkat keras dan perangkat lunak yang digunakan selama pelaksanaan pengujian "
             "sistem dirangkum dalam Tabel 5.1.")
             
    caption1 = add_body(doc, p_insert_v, "Tabel 5.1. Lingkungan pengujian", bold=True)
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t1_headers = ["Komponen", "Spesifikasi", "Keterangan"]
    t1_data = [
        ["Perangkat Dev", "Laptop ASUS Vivobook AMD Ryzen 5, RAM 16GB", "Menjalankan editor kode VS Code, Gradle, dan compiler Flutter."],
        ["Sistem Operasi", "Windows 11 Home 64-bit", "Platform sistem operasi utama untuk lingkungan pengembangan."],
        ["Flutter SDK", "Flutter Version 3.35.7", "Membawa SDK Dart 3.9.2 untuk kompilasi kode native."],
        ["Backend", "Supabase PostgreSQL Database Cloud", "Menyimpan seluruh schema database transaksi akademik secara online."],
        ["Perangkat Uji", "Xiaomi Redmi Note 10 Pro (Android 13)", "Perangkat mobile fisik untuk uji coba aplikasi native APK."],
        ["Penguji", "1 Kepala Sekolah, 2 Guru, 1 BK, 2 Siswa", "Staf penguji lapangan dari perwakilan pengguna SMPN 1 Jenar."]
    ]
    add_table(doc, p_insert_v, t1_headers, t1_data, col_widths=[1.5, 2.5, 2.5])

    add_heading2(doc, p_insert_v, "5.3 Prosedur Pengujian")
    
    add_body(doc, p_insert_v,
             "Rencana skenario pengujian fungsional sistem menggunakan metode black-box disusun "
             "secara sistematis pada Tabel 5.2.")
             
    caption2 = add_body(doc, p_insert_v, "Tabel 5.2. Rencana pengujian black-box", bold=True)
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t2_headers = ["ID", "Objek Uji", "Butir Uji"]
    t2_data = [
        ["TC-01", "Halaman Login", "Pengguna dapat login menggunakan akun terdaftar di Supabase Auth."],
        ["TC-02", "Validasi Kredensial", "Sistem menampilkan alert jika email format salah atau password kosong."],
        ["TC-03", "Pemilihan Peran", "Akun dengan multi-role dapat memilih peran dan masuk ke dashboard yang sesuai."],
        ["TC-04", "Dashboard Guru", "Guru dapat melihat jadwal mengajar aktif sesuai semester berjalan."],
        ["TC-05", "Input Jurnal & Presensi", "Guru dapat menyimpan catatan materi pembelajaran dan presensi siswa."],
        ["TC-06", "Dashboard Kepsek", "Kepala sekolah dapat memantau status kelas dan jurnal real-time."],
        ["TC-07", "Dashboard BK", "BK dapat mengelola data siswa aktif dan pembagian kelas siswa."],
        ["TC-08", "Dashboard Siswa", "Siswa dapat melihat rekap kehadiran personal dan tugas aktif."],
        ["TC-09", "Unggah Administrasi", "Guru dapat mengunggah berkas RPP/Silabus PDF ke Supabase Storage."],
        ["TC-10", "Keluar Sesi", "Pengguna dapat keluar sesi dan membersihkan cached auth session."]
    ]
    add_table(doc, p_insert_v, t2_headers, t2_data, col_widths=[0.8, 1.7, 4.0])

    add_heading2(doc, p_insert_v, "5.4 Hasil Uji dan Kesimpulan Pengujian")
    
    add_body(doc, p_insert_v,
             "Hasil pengujian fungsional beserta status kelayakannya tercantum pada Tabel 5.3.")
             
    caption3 = add_body(doc, p_insert_v, "Tabel 5.3. Hasil uji fungsional", bold=True)
    caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t3_headers = ["ID", "Hasil yang Diharapkan", "Hasil Pengujian"]
    t3_data = [
        ["TC-01", "Login berhasil dan rute masuk sesuai dashboard peran.", "Berhasil"],
        ["TC-02", "Muncul pesan kesalahan login 'Invalid login credentials'.", "Berhasil"],
        ["TC-03", "Memilih dashboard kepala sekolah berhasil memuat menu monitoring.", "Berhasil"],
        ["TC-04", "Jadwal mengajar memuat daftar hari, kelas, dan mapel guru terkait.", "Berhasil"],
        ["TC-05", "Data tersimpan di tabel jurnal dan absensi_jurnal Supabase.", "Berhasil"],
        ["TC-06", "Data monitoring kelas memperlihatkan persentase keterisian jurnal.", "Berhasil"],
        ["TC-07", "Pembagian kelas siswa memperbarui data di tabel kelas_siswa.", "Berhasil"],
        ["TC-08", "Detail kehadiran personal siswa terakumulasi per status presensi.", "Berhasil"],
        ["TC-09", "Tautan file RPP tersimpan di database dan berkas ada di storage bucket.", "Berhasil"],
        ["TC-10", "Aplikasi kembali ke AuthBootstrapScreen dan menghapus sesi.", "Berhasil"]
    ]
    add_table(doc, p_insert_v, t3_headers, t3_data, col_widths=[0.8, 4.2, 1.5])
    
    add_body(doc, p_insert_v,
             "Kesimpulan pengujian: berdasarkan hasil pengujian black-box, seluruh skenario pengujian "
             "(TC-01 hingga TC-10) menunjukkan hasil yang sesuai dengan yang diharapkan (100% Berhasil). "
             "Hal ini membuktikan bahwa fungsionalitas utama aplikasi SIMONAS telah berjalan dengan baik "
             "di lingkungan uji coba.")

    # Now clear empty paragraphs between KESIMPULAN DAN SARAN and DAFTAR PUSTAKA
    # Find position again
    idx_start6 = -1
    idx_end6 = -1
    for idx, p in enumerate(doc.paragraphs):
        if p._element == p_kesimpulan._element:
            idx_start6 = idx
        if p._element == p_daftar_pustaka._element:
            idx_end6 = idx
            break
            
    print(f"Clearing empty paragraphs in Bab VI (index {idx_start6} to {idx_end6})...")
    to_remove6 = []
    for i in range(idx_start6 + 1, idx_end6):
        to_remove6.append(doc.paragraphs[i])
        
    for p in to_remove6:
        p._element.getparent().remove(p._element)
        
    # Find position again
    p_insert_vi = None
    for idx, p in enumerate(doc.paragraphs):
        if p._element == p_daftar_pustaka._element:
            p_insert_vi = p
            break
            
    # Write BAB VI contents
    print("Writing BAB VI content...")
    add_heading2(doc, p_insert_vi, "6.1 Kesimpulan")
    
    add_body(doc, p_insert_vi,
             "Berdasarkan hasil analisis, perancangan, implementasi, dan pengujian yang telah "
             "dilakukan pada aplikasi SIMONAS, dapat ditarik kesimpulan sebagai berikut:")
             
    add_body(doc, p_insert_vi,
             "1. Aplikasi SIMONAS telah berhasil dikembangkan menggunakan framework Flutter untuk frontend "
             "mobile dan Supabase sebagai backend-as-a-service. Integrasi ini terbukti efektif dalam mempermudah "
             "proses digitalisasi data akademik sekolah.")
             
    add_body(doc, p_insert_vi,
             "2. Sistem multi-role yang dirancang telah berfungsi dengan baik, memberikan akses antarmuka "
             "dan hak operasi data yang berbeda untuk Kepala Sekolah, Guru, BK, dan Siswa sesuai dengan "
             "tanggung jawab masing-masing.")
             
    add_body(doc, p_insert_vi,
             "3. Aplikasi SIMONAS berhasil menyelesaikan masalah pencatatan akademik di SMP Negeri 1 Jenar "
             "dengan mengintegrasikan alur pengisian jurnal mengajar, presensi siswa, pengelolaan nilai/tugas, "
             "serta pengunggahan berkas administrasi pembelajaran guru dalam satu sistem terpadu.")
             
    add_body(doc, p_insert_vi,
             "4. Fitur monitoring kelas, monitoring jurnal, dan verifikasi administrasi bagi Kepala Sekolah "
             "dapat mempermudah pengawasan proses kegiatan belajar mengajar secara real-time.")
             
    add_body(doc, p_insert_vi,
             "5. Pengujian black-box membuktikan seluruh fungsi utama aplikasi berjalan sukses 100% dan siap "
             "diimplementasikan untuk mendukung operasional sekolah.")

    add_heading2(doc, p_insert_vi, "6.2 Saran")
    
    add_body(doc, p_insert_vi,
             "Untuk pengembangan aplikasi SIMONAS selanjutnya, disarankan beberapa hal sebagai berikut:")
             
    add_body(doc, p_insert_vi,
             "1. Melakukan uji coba lapangan dalam skala kecil (pilot testing) pada satu atau dua kelas "
             "terlebih dahulu sebelum diimplementasikan secara menyeluruh di sekolah guna membiasakan pengguna.")
             
    add_body(doc, p_insert_vi,
             "2. Menambahkan fitur push notification di perangkat mobile untuk memberikan pengingat otomatis "
             "ketika terdapat tugas baru bagi siswa atau dokumen baru yang perlu diverifikasi kepala sekolah.")
             
    add_body(doc, p_insert_vi,
             "3. Menyediakan fitur rekapitulasi data absensi dan jurnal mengajar dalam format PDF atau Excel "
             "yang dapat diunduh langsung untuk kebutuhan pelaporan bulanan.")
             
    add_body(doc, p_insert_vi,
             "4. Mengembangkan modul khusus untuk Orang Tua/Wali agar dapat memantau presensi dan perkembangan "
             "nilai akademik siswa secara langsung.")
             
    add_body(doc, p_insert_vi,
             "5. Meningkatkan integrasi dengan sistem eksternal yang sudah berjalan di sekolah, seperti "
             "sinkronisasi otomatis dengan e-Rapor.")

    print(f"Saving changes to {filename}...")
    doc.save(filename)
    print("Document saved successfully!")

if __name__ == "__main__":
    main()

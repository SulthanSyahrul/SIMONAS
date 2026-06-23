import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    
    targets = [94, 100, 115, 157]
    for p_idx in targets:
        print(f"\n=== Paragraph {p_idx} Runs ===")
        p = doc.paragraphs[p_idx]
        print(f"Full Text: {p.text}")
        for r_idx, run in enumerate(p.runs):
            print(f"  Run {r_idx}: text='{run.text}', bold={run.bold}, italic={run.italic}, font='{run.font.name}', size={run.font.size}")

    print("\n=== Table 3.1 (Index 8) Cells ===")
    t = doc.tables[8]
    for r_idx, row in enumerate(t.rows):
        print(f"Row {r_idx}:")
        for c_idx, cell in enumerate(row.cells):
            print(f"  Col {c_idx} text: '{cell.text}'")
            for p_idx, p in enumerate(cell.paragraphs):
                for r_idx2, run in enumerate(p.runs):
                    print(f"    P{p_idx} Run{r_idx2}: text='{run.text}', bold={run.bold}, italic={run.italic}")

if __name__ == '__main__':
    main()

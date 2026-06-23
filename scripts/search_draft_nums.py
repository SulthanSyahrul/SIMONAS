import docx

doc = docx.Document('Draft Template Naskah TA 2025.docx')
for idx, p in enumerate(doc.paragraphs):
    if '6.3' in p.text or '5.3' in p.text or '6.4' in p.text or '5.4' in p.text:
        print(f"P#{idx}: {p.text}")

for idx, t in enumerate(doc.tables):
    for r_idx, r in enumerate(t.rows):
        for c_idx, cell in enumerate(r.cells):
            if '6.3' in cell.text or '5.3' in cell.text or '6.4' in cell.text or '5.4' in cell.text:
                print(f"T#{idx} R{r_idx} C{c_idx}: {cell.text[:100]}")

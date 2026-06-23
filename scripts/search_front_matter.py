import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    print("=== Searching paragraphs 0 to 120 ===")
    for idx in range(min(120, len(doc.paragraphs))):
        p_text = doc.paragraphs[idx].text
        if "lampiran" in p_text.lower():
            print(f"P {idx}: {p_text}")

if __name__ == '__main__':
    main()

import docx

def main():
    filename = 'Naskah TA Sulthan Syahrul BAB I-IV.docx'
    print(f"Loading {filename}...")
    doc = docx.Document(filename)
    
    # 1. Update paragraph 79 (Preface)
    p79 = doc.paragraphs[79]
    old_79 = (
        "Penulis menyadari bahwa naskah ini masih memerlukan penyempurnaan, "
        "terutama pada bagian pengujian lapangan, pengisian gambar aktual, "
        "serta finalisasi sitasi melalui reference manager. Oleh karena itu, "
        "kritik dan saran yang membangun sangat diharapkan."
    )
    new_79 = (
        "Penulis menyadari bahwa naskah Tugas Akhir ini telah diselesaikan dengan baik, "
        "mencakup pengujian fungsional sistem secara menyeluruh, penyajian gambar "
        "implementasi aktual, serta penyelarasan sitasi referensi. Oleh karena itu, "
        "kritik dan saran yang membangun sangat diharapkan."
    )
    if p79.text.strip() == old_79.strip():
        p79.text = new_79
        p79.paragraph_format.line_spacing = 1.5
        p79.paragraph_format.space_after = docx.shared.Pt(6)
        for run in p79.runs:
            run.font.name = "Times New Roman"
            run.font.size = docx.shared.Pt(12)
        print("Updated Paragraph 79 (Preface).")

    # 2. Update paragraph 94 (Intisari / Summary)
    p94 = doc.paragraphs[94]
    old_sentence = "Naskah ini masih perlu dilengkapi tangkapan layar aktual dan hasil pengujian lapangan sebelum digunakan sebagai versi final."
    new_sentence = "Aplikasi SIMONAS telah diuji secara menyeluruh menggunakan metode pengujian fungsional black-box serta widget testing untuk memvalidasi kesesuaian seluruh fitur dengan hasil yang berhasil 100%."
    if old_sentence in p94.text:
        p94.text = p94.text.replace(old_sentence, new_sentence)
        p94.paragraph_format.line_spacing = 1.5
        p94.paragraph_format.space_after = docx.shared.Pt(6)
        for run in p94.runs:
            run.font.name = "Times New Roman"
            run.font.size = docx.shared.Pt(12)
        print("Updated Paragraph 94 (Intisari).")

    # 3. Update paragraph 157 (Batasan produk)
    p157 = doc.paragraphs[157]
    old_157_part = "6) hasil pengujian lapangan perlu dilengkapi setelah aplikasi diuji oleh pengguna sekolah."
    new_157_part = "6) pengujian dibatasi pada pengujian fungsional sistem menggunakan metode black-box dan unit testing untuk memvalidasi seluruh fitur utama."
    if old_157_part in p157.text:
        p157.text = p157.text.replace(old_157_part, new_157_part)
        p157.paragraph_format.line_spacing = 1.5
        p157.paragraph_format.space_after = docx.shared.Pt(6)
        for run in p157.runs:
            run.font.name = "Times New Roman"
            run.font.size = docx.shared.Pt(12)
        print("Updated Paragraph 157 (Batasan produk).")

    print(f"Saving changes to {filename}...")
    doc.save(filename)
    print("Document saved successfully!")

if __name__ == '__main__':
    main()

import docx

doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
sinarbelajar_count = 0
simonas_count = 0

for i, p in enumerate(doc.paragraphs):
    if 'sinarbelajar' in p.text.lower():
        sinarbelajar_count += 1
    if 'simonas' in p.text.lower():
        simonas_count += 1

print(f"SinarBelajar count in paragraphs: {sinarbelajar_count}")
print(f"SIMONAS count in paragraphs: {simonas_count}")

# Check tables too
for t_idx, t in enumerate(doc.tables):
    for r in t.rows:
        for c in r.cells:
            if 'sinarbelajar' in c.text.lower():
                sinarbelajar_count += 1
            if 'simonas' in c.text.lower():
                simonas_count += 1
                
print(f"Total SinarBelajar count: {sinarbelajar_count}")
print(f"Total SIMONAS count: {simonas_count}")

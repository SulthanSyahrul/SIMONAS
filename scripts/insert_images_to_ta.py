import os
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_images():
    doc_path = "Naskah_TA_SinarBelajar_SMPN1Jenar_v3_cited.docx"
    print(f"Opening {doc_path}...")
    doc = docx.Document(doc_path)
    
    # Mapping of target paragraph text to (image_path, width_in_inches)
    mapping = {
        "Gambar 1.1. Kondisi proses akademik sebelum digitalisasi": ("scripts/extracted_images/diagram_manual_process.png", 5.8),
        "Gambar 2.1. Gambaran umum produk SinarBelajar": ("scripts/extracted_images/diagram_sinarbelajar_overview.png", 5.8),
        "Gambar 3.1. Arsitektur sistem client-server": ("scripts/extracted_images/image13.png", 5.8),
        "Gambar 3.2. Use case diagram aplikasi": ("scripts/extracted_images/image3.png", 5.5),
        "Gambar 3.3. Activity diagram login dan pemilihan role": ("scripts/extracted_images/diagram_activity_login.png", 5.5),
        "Gambar 3.4. Activity diagram input jurnal dan absensi": ("scripts/extracted_images/diagram_activity_jurnal.png", 5.5),
        "Gambar 3.5. ERD Modul Pengguna": ("scripts/extracted_images/image9.png", 5.8),
        "Gambar 3.6. ERD Modul Jadwal dan Jurnal": ("scripts/extracted_images/image10.png", 5.8),
        "Gambar 3.7. ERD Modul Penilaian": ("scripts/extracted_images/image11.png", 5.8),
        "Gambar 3.8. ERD Modul Administrasi Pembelajaran": ("scripts/extracted_images/image12.png", 5.8),
        "Lampiran E.1. ERD Modul Pengelolaan Kelas": ("scripts/extracted_images/image17.png", 5.8),
        "Lampiran E.2. ERD Modul Histori Wali Kelas": ("scripts/extracted_images/image18.png", 5.8)
    }
    
    # We iterate backwards or keep track of indices, but wait!
    # Inserting a paragraph before paragraph i does not affect the text of the following paragraphs, 
    # but it increases doc.paragraphs length. So if we search by text, it's safer to find matching paragraphs,
    # store them in a list, and then process them.
    
    to_modify = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in mapping:
            to_modify.append((p, text))
            
    print(f"Found {len(to_modify)} target paragraphs to insert images before.")
    
    for p, text in to_modify:
        img_path, width_inch = mapping[text]
        if not os.path.exists(img_path):
            print(f"Warning: Image file not found: {img_path}")
            continue
            
        print(f"Inserting {img_path} before paragraph: '{text}'")
        # Insert paragraph before the caption
        new_p = p.insert_paragraph_before()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the image to this paragraph
        run = new_p.add_run()
        run.add_picture(img_path, width=Inches(width_inch))
        
        # Add an empty paragraph after the image (spacing)
        spacer = p.insert_paragraph_before()
        spacer.paragraph_format.space_after = docx.shared.Pt(6)
        
    out_path = "Naskah_TA_SinarBelajar_SMPN1Jenar_v3_cited_filled.docx"
    try:
        doc.save(doc_path)
        print(f"Document saved successfully as {doc_path}")
        try:
            import shutil
            shutil.copy(doc_path, out_path)
            print(f"Copied updated document to {out_path}")
        except PermissionError:
            print(f"Warning: {out_path} is locked, could not copy.")
    except PermissionError:
        print(f"Warning: {doc_path} is locked (likely open in Word).")
        try:
            doc.save(out_path)
            print(f"Saved filled version as new file: {out_path}")
        except PermissionError:
            print(f"Error: Both {doc_path} and {out_path} are locked.")

if __name__ == "__main__":
    insert_images()


import docx
import os

def inspect_file(filename, output_txt):
    if not os.path.exists(filename):
        print(f"File {filename} does not exist.")
        return
    
    print(f"Inspecting {filename}...")
    doc = docx.Document(filename)
    
    with open(output_txt, 'w', encoding='utf-8') as f:
        # Write total paragraphs and tables count
        f.write(f"=== File: {filename} ===\n")
        f.write(f"Paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Tables: {len(doc.tables)}\n\n")
        
        # Write headings and sections
        f.write("=== Headings / Outline ===\n")
        for idx, p in enumerate(doc.paragraphs):
            # Check if paragraph looks like a heading
            text_strip = p.text.strip()
            if p.style.name.startswith('Heading') or text_strip.upper().startswith('BAB ') or (len(text_strip) < 100 and any(keyword in text_strip.upper() for keyword in ['KATA PENGANTAR', 'DAFTAR ISI', 'DAFTAR GAMBAR', 'DAFTAR TABEL', 'INTISARI', 'ABSTRACT', 'KESIMPULAN', 'SARAN'])):
                f.write(f"[P#{idx} Style={p.style.name}]: '{text_strip}'\n")
        
        # Also let's inspect where BAB V and BAB VI are located and print their surrounding paragraphs
        f.write("\n=== BAB V & VI Content and surrounding paragraphs ===\n")
        found_bab5 = False
        found_bab6 = False
        
        for idx, p in enumerate(doc.paragraphs):
            text_strip = p.text.strip()
            if 'BAB V' in text_strip.upper():
                f.write(f"\n--- Found BAB V at Paragraph {idx} ---\n")
                found_bab5 = True
                # print 30 paragraphs after this
                for j in range(max(0, idx - 2), min(len(doc.paragraphs), idx + 100)):
                    f.write(f"[{j}]: {doc.paragraphs[j].text}\n")
            if 'BAB VI' in text_strip.upper():
                f.write(f"\n--- Found BAB VI at Paragraph {idx} ---\n")
                found_bab6 = True
                for j in range(max(0, idx - 2), min(len(doc.paragraphs), idx + 100)):
                    f.write(f"[{j}]: {doc.paragraphs[j].text}\n")
                    
        if not found_bab5:
            f.write("\nBAB V not found by direct match. Searching case-insensitive...\n")
            for idx, p in enumerate(doc.paragraphs):
                if 'bab v' in p.text.lower():
                    f.write(f"Potential BAB V match at P#{idx}: '{p.text}'\n")
        if not found_bab6:
            f.write("\nBAB VI not found by direct match. Searching case-insensitive...\n")
            for idx, p in enumerate(doc.paragraphs):
                if 'bab vi' in p.text.lower():
                    f.write(f"Potential BAB VI match at P#{idx}: '{p.text}'\n")

# Run on the files
inspect_file('Draft Template Naskah TA 2025.docx', 'scripts/inspect_draft_results.txt')
inspect_file('Naskah TA Sulthan Syahrul BAB I-IV.docx', 'scripts/inspect_ta_sultan_results.txt')
print("Inspection complete.")

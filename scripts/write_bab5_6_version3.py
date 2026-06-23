import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import copy

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)

def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.5):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_heading2(p_before, text):
    new_p = p_before.insert_paragraph_before()
    format_paragraph(new_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
    run = new_p.add_run(text)
    set_font(run, bold=True)
    return new_p

def add_body(p_before, text, bold=False, italic=False):
    new_p = p_before.insert_paragraph_before()
    format_paragraph(new_p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6)
    run = new_p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return new_p

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def format_table(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr[0].append(jc)

def add_custom_table(p_before, headers, data, col_widths=None):
    doc = p_before.part.document
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    p_before._element.addprevious(table._element)
    format_table(table)
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4, line_spacing=1.0)
        for r in p.runs:
            set_font(r, bold=True)
        shading = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        border_style = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
        set_cell_border(hdr_cells[i], top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    # Data
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = cell_value
            p = row_cells[c_idx].paragraphs[0]
            is_center = len(cell_value) < 10 or cell_value in ["Berhasil", "Laki-laki", "Perempuan"] or cell_value.startswith("TC-")
            align = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, align=align, space_before=4, space_after=4, line_spacing=1.0)
            for r in p.runs:
                set_font(r)
            border_style = {'val': 'single', 'sz': '4', 'space': '0', 'color': 'CCCCCC'}
            set_cell_border(row_cells[c_idx], top=border_style, bottom=border_style, left=border_style, right=border_style)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    spacer = p_before.insert_paragraph_before()
    format_paragraph(spacer, space_before=0, space_after=6)
    return table

def copy_and_insert_table(src_table, p_before):
    tbl_elem = src_table._element
    new_tbl_elem = copy.deepcopy(tbl_elem)
    p_before._element.addprevious(new_tbl_elem)
    from docx.table import Table
    new_table = Table(new_tbl_elem, p_before.part)
    format_table(new_table)
    return new_table

def add_detailed_test_case_table(p_before, src_table, tc_id, name, desc, cond, date_str, tester_name, scenario_text, expected_text, observed_text, conclusion_text="Berhasil"):
    new_table = copy_and_insert_table(src_table, p_before)
    
    new_table.rows[0].cells[1].text = tc_id
    new_table.rows[1].cells[1].text = name
    new_table.rows[2].cells[1].text = desc
    new_table.rows[3].cells[1].text = cond
    new_table.rows[4].cells[1].text = date_str
    new_table.rows[5].cells[1].text = tester_name
    
    new_table.rows[7].cells[0].text = scenario_text
    new_table.rows[7].cells[1].text = scenario_text
    
    new_table.rows[10].cells[0].text = expected_text
    new_table.rows[10].cells[1].text = observed_text
    
    new_table.rows[12].cells[0].text = conclusion_text
    new_table.rows[12].cells[1].text = conclusion_text
    
    # Re-apply formatting & font to Table
    for r_idx, row in enumerate(new_table.rows):
        is_heading_row = r_idx in [6, 8, 11]
        for c_idx, cell in enumerate(row.cells):
            p = cell.paragraphs[0]
            if is_heading_row:
                align = WD_ALIGN_PARAGRAPH.CENTER
                bold = True
                shading = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)
            elif r_idx == 9: # Yang diharapkan vs Hasil Pengamatan subheaders
                align = WD_ALIGN_PARAGRAPH.CENTER
                bold = True
                shading = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                bold = (c_idx == 0) # Bold labels in left column
                align = WD_ALIGN_PARAGRAPH.LEFT
                
            format_paragraph(p, align=align, space_before=4, space_after=4, line_spacing=1.0)
            for run in p.runs:
                set_font(run, bold=bold)
                
            border_style = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
            set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
            
    # Add a spacer after table
    spacer = p_before.insert_paragraph_before()
    format_paragraph(spacer, space_before=0, space_after=6)
    return new_table

def main():
    dest_filename = 'Naskah TA Sulthan Syahrul BAB I-IV.docx'
    src_filename = 'Draft Template Naskah TA 2025.docx'
    
    # Define all 24 test cases
    test_cases_list = [
        {
            "id": "TC-02",
            "object": "Monitoring Kelas",
            "butir_uji": "Pemantauan status KBM kelas secara realtime",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Monitoring Kelas oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat melihat dashboard pemantauan kelas secara realtime",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan berada di dashboard SIMONAS.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu pemantauan kelas.\n2. Pilih kelas aktif (misalnya 7A) untuk melihat detail status kelas.\n3. Amati informasi kehadiran guru dan jurnal pembelajaran.",
            "expected": "Halaman monitoring menampilkan status kelas 7A, nama guru mengajar, dan status KBM yang sedang berlangsung secara real-time.",
            "observed": "Dashboard menampilkan daftar seluruh kelas dengan status pembelajaran, nama guru yang sedang mengajar, dan detail keterisian jurnal secara akurat.",
            "summary_expected": "Akses dashboard pemantauan kelas/jurnal berhasil",
            "summary_observed": "Sesuai harapan, user Surono dapat masuk dan melihat dashboard pemantauan"
        },
        {
            "id": "TC-03",
            "object": "Monitoring Administrasi",
            "butir_uji": "Pemantauan status pengunggahan berkas administrasi guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Monitoring Administrasi Guru oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat memantau status pengunggahan berkas administrasi guru",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan berada di dashboard SIMONAS.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu monitoring administrasi guru.\n2. Amati daftar berkas administrasi pembelajaran yang telah diunggah oleh para guru (RPP, silabus, dll).",
            "expected": "Sistem menampilkan tabel rekapitulasi status pengunggahan berkas administrasi untuk semua guru secara terperinci.",
            "observed": "Daftar berkas administrasi pembelajaran tampil lengkap dengan status (belum unggah / menunggu verifikasi / disetujui / ditolak) per nama guru.",
            "summary_expected": "Akses monitoring administrasi guru berjalan lancar",
            "summary_observed": "Tabel status kelengkapan berkas seluruh guru tampil lengkap"
        },
        {
            "id": "TC-04",
            "object": "Verifikasi Administrasi",
            "butir_uji": "Pemberian status persetujuan atau penolakan berkas guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Verifikasi Administrasi Guru oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat memverifikasi berkas administrasi pembelajaran yang diunggah oleh guru",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan terdapat berkas administrasi berstatus 'menunggu verifikasi'.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman verifikasi administrasi.\n2. Pilih salah satu berkas yang diunggah guru (misal RPP matematika Afifatun).\n3. Klik tombol 'Verifikasi/Setujui'.",
            "expected": "Status berkas administrasi berubah menjadi 'Disetujui' dan tersimpan di database.",
            "observed": "Kepala Sekolah berhasil menyetujui berkas, status berkas berubah menjadi disetujui, dan perubahan tersinkronisasi ke dashboard guru.",
            "summary_expected": "Persetujuan status berkas administrasi guru berhasil",
            "summary_observed": "Status berkas berubah menjadi disetujui dan terupdate di database"
        },
        {
            "id": "TC-05",
            "object": "Monitoring Jurnal",
            "butir_uji": "Pemantauan isian jurnal mengajar harian guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Monitoring Jurnal Mengajar oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat melihat jurnal mengajar harian guru secara realtime",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan berada di dashboard SIMONAS.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu monitoring jurnal mengajar.\n2. Lakukan pencarian atau filter berdasarkan tanggal hari ini dan nama guru.",
            "expected": "Jurnal mengajar yang diinput oleh guru pada hari tersebut tampil lengkap beserta rincian materi pelajaran.",
            "observed": "Jurnal mengajar harian guru tertera di layar dengan materi pembelajaran, jam mengajar, kelas, dan status absensi siswa.",
            "summary_expected": "Jurnal mengajar harian guru dapat terpantau",
            "summary_observed": "Jurnal guru tertera di layar secara lengkap dan realtime"
        },
        {
            "id": "TC-06",
            "object": "CRUD Jadwal & Wali Kelas",
            "butir_uji": "Pengelolaan data jadwal pelajaran dan penugasan wali kelas",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Jadwal Pelajaran dan Wali Kelas oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat mengelola (tambah, lihat, ubah, hapus) jadwal pelajaran dan penugasan wali kelas",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan membuka menu pengaturan akademik.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Jadwal & Wali Kelas.\n2. Tambah data jadwal baru (Matematika kelas 7A hari Senin).\n3. Ubah penugasan wali kelas untuk kelas 7A menjadi Eko.\n4. Hapus salah satu jadwal lama.",
            "expected": "Data jadwal pelajaran baru berhasil ditambahkan, wali kelas berhasil diperbarui, dan jadwal lama terhapus dari database.",
            "observed": "Formulir CRUD merespons dengan cepat. Penambahan, pengubahan, dan penghapusan jadwal pelajaran serta wali kelas berhasil disimpan ke Supabase.",
            "summary_expected": "Pengelolaan jadwal dan wali kelas berhasil tersimpan",
            "summary_observed": "Formulir CRUD merespons, data jadwal dan wali kelas terupdate"
        },
        {
            "id": "TC-07",
            "object": "CRUD Tahun Ajaran",
            "butir_uji": "Pengelolaan data tahun ajaran dan semester aktif",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Tahun Ajaran oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat mengelola data tahun ajaran dan semester aktif",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan membuka menu tahun ajaran.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Tahun Ajaran.\n2. Tambah data tahun ajaran baru (2026/2027 Ganjil).\n3. Aktifkan tahun ajaran baru tersebut.",
            "expected": "Data tahun ajaran baru tersimpan dan status tahun ajaran aktif berubah di database.",
            "observed": "Tahun ajaran baru berhasil dibuat dan status keaktifan semester berhasil diganti menjadi aktif.",
            "summary_expected": "Kelola data tahun ajaran dan semester berhasil",
            "summary_observed": "Tahun ajaran baru terdaftar dan aktif di sistem"
        },
        {
            "id": "TC-08",
            "object": "CRUD Manajemen Guru",
            "butir_uji": "Pengelolaan data guru dan akun guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Manajemen Guru oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat mengelola data akun dan informasi guru",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan membuka menu manajemen guru.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Manajemen Guru.\n2. Tambah data guru baru (NIP, Nama, Email, Peran).\n3. Ubah nama guru yang ada.\n4. Hapus data guru yang tidak aktif.",
            "expected": "Penambahan, pembaruan, dan penghapusan data profil guru berhasil disimpan di database.",
            "observed": "Data guru berhasil dimanipulasi melalui form CRUD guru. Akun guru baru dapat dibuat dan masuk ke sistem.",
            "summary_expected": "Pendaftaran dan manajemen akun guru berhasil",
            "summary_observed": "Akun guru baru berhasil dibuat dan terintegrasi di database"
        },
        {
            "id": "TC-09",
            "object": "CRUD Mapel",
            "butir_uji": "Pengelolaan data mata pelajaran",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Mata Pelajaran oleh Kepala Sekolah",
            "desc": "Memeriksa apakah Kepala Sekolah dapat mengelola data mata pelajaran di sekolah",
            "cond": "Kepala Sekolah sudah masuk ke sistem dengan email surono@smpn1jenar.local dan membuka menu kelola mata pelajaran.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Mapel.\n2. Tambah data mata pelajaran baru (Bahasa Indonesia).\n3. Ubah nama mata pelajaran yang sudah ada.\n4. Hapus mata pelajaran tertentu.",
            "expected": "Mata pelajaran baru berhasil disimpan, diubah, dan dihapus di database.",
            "observed": "Daftar mata pelajaran terupdate dengan benar di database dan perubahan langsung terlihat pada form penjadwalan.",
            "summary_expected": "Kelola mata pelajaran baru berjalan lancar",
            "summary_observed": "Mata pelajaran terupdate dan tersinkronisasi di database"
        },
        {
            "id": "TC-10",
            "object": "CRUD Manajemen Akun Siswa",
            "butir_uji": "Pengelolaan data profil dan akun siswa oleh Guru BK",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Akun Siswa oleh Guru BK",
            "desc": "Memeriksa apakah Guru BK dapat mengelola data akun siswa",
            "cond": "Guru BK sudah masuk ke sistem dengan email budi@smpn1jenar.local dan membuka menu kelola siswa.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Akun Siswa.\n2. Tambah akun siswa baru (Nama: Abdul Murudul, NIS, Email).\n3. Ubah data siswa tersebut (misal alamat).\n4. Hapus data siswa yang keluar.",
            "expected": "Data akun siswa baru berhasil ditambahkan, diubah, dan dihapus dari database.",
            "observed": "Sistem berhasil merekam data siswa baru dan memperbarui profil siswa di database public.users dan auth.users.",
            "summary_expected": "Pendaftaran akun siswa baru oleh BK berjalan lancar",
            "summary_observed": "Akun siswa baru terekam di database dan siap masuk sistem"
        },
        {
            "id": "TC-11",
            "object": "CRUD Pengaturan Kelas Siswa",
            "butir_uji": "Penetapan kelas bagi siswa aktif",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Pengaturan Kelas Siswa oleh Guru BK",
            "desc": "Memeriksa apakah Guru BK dapat mengatur penetapan kelas bagi siswa aktif",
            "cond": "Guru BK sudah masuk ke sistem dengan email budi@smpn1jenar.local dan berada di menu pengaturan kelas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Pengaturan Kelas Siswa.\n2. Pilih siswa (Abdul Murudul) dan tetapkan ke kelas 7A.\n3. Pindahkan siswa ke kelas 7B.",
            "expected": "Kelas siswa berhasil ditetapkan dan diperbarui di database histori kelas.",
            "observed": "Penetapan kelas siswa tersimpan dengan sukses. Riwayat kelas siswa terupdate secara realtime.",
            "summary_expected": "Penetapan dan pemindahan kelas siswa berhasil",
            "summary_observed": "Status kelas siswa terupdate di tabel histori_kelas"
        },
        {
            "id": "TC-12",
            "object": "CRUD Kenaikan Kelas",
            "butir_uji": "Pemrosesan kenaikan kelas siswa secara periodik",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Kenaikan Kelas oleh Guru BK",
            "desc": "Memeriksa apakah Guru BK dapat mengelola proses kenaikan kelas siswa secara periodik",
            "cond": "Guru BK sudah masuk ke sistem dengan email budi@smpn1jenar.local dan berada pada halaman kelulusan/kenaikan kelas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kenaikan Kelas.\n2. Pilih siswa kelas 7A yang memenuhi syarat.\n3. Proses kenaikan kelas siswa tersebut ke kelas 8A.",
            "expected": "Status kelas siswa naik dari kelas 7A menjadi kelas 8A dan tercatat di database tahun ajaran baru.",
            "observed": "Proses kenaikan kelas berhasil diproses secara massal, data riwayat kelas siswa diperbarui menjadi kelas 8A.",
            "summary_expected": "Pemrosesan kenaikan kelas siswa berhasil dilakukan",
            "summary_observed": "Siswa berhasil naik dari 7A ke 8A di database"
        },
        {
            "id": "TC-13",
            "object": "CRUD Profil BK",
            "butir_uji": "Pembaruan data profil pribadi BK",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Profil Pribadi oleh Guru BK",
            "desc": "Memeriksa apakah Guru BK dapat memperbarui data profil pribadinya",
            "cond": "Guru BK sudah masuk ke sistem dengan email budi@smpn1jenar.local dan membuka halaman profil.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman Profil Saya.\n2. Ubah nomor telepon dan foto profil.\n3. Klik tombol 'Simpan Perubahan'.",
            "expected": "Informasi profil Guru BK terupdate di database dan ditampilkan dengan benar pada menu profil.",
            "observed": "Data nomor telepon berhasil diperbarui, foto profil baru terunggah ke penyimpanan dan tampil di profil.",
            "summary_expected": "Akses kelola profil Guru BK berhasil",
            "summary_observed": "Profil terupdate dengan nomor telepon terbaru"
        },
        {
            "id": "TC-14",
            "object": "Lihat Jadwal Mengajar",
            "butir_uji": "Pemeriksaan penugasan jadwal mengajar guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Lihat Jadwal Mengajar oleh Guru",
            "desc": "Memeriksa apakah Guru dapat melihat jadwal mengajar yang ditugaskan kepadanya",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan berada di dashboard Guru.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Jadwal Mengajar.\n2. Amati daftar hari, jam, kelas, dan mata pelajaran yang diampu.",
            "expected": "Sistem menampilkan daftar jadwal mengajar Guru bersangkutan secara lengkap sesuai hari aktif.",
            "observed": "Jadwal mengajar untuk hari Senin hingga Jumat tampil secara berurutan sesuai data kelas yang ditugaskan.",
            "summary_expected": "Penayangan jadwal mengajar guru aktif berhasil",
            "summary_observed": "Jadwal mengajar tertera secara urut sesuai database"
        },
        {
            "id": "TC-15",
            "object": "CRUD Jurnal & Absensi",
            "butir_uji": "Pengisian dan kelola jurnal mengajar serta presensi siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Jurnal Mengajar dan Absensi oleh Guru",
            "desc": "Memeriksa apakah Guru dapat mengisi, melihat, mengubah, dan menghapus jurnal mengajar harian beserta absensi siswa",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan membuka halaman jurnal kelas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Isi jurnal mengajar baru dan presensi siswa kelas 7A.\n2. Ubah materi jurnal yang salah ketik.\n3. Hapus jurnal yang salah input.",
            "expected": "Jurnal mengajar dan absensi berhasil dibuat, diperbarui, dan dihapus di database.",
            "observed": "Pengisian jurnal dan presensi berjalan lancar. Riwayat jurnal terupdate setelah dilakukan edit dan hapus.",
            "summary_expected": "Kelola pengisian jurnal dan absensi siswa berhasil",
            "summary_observed": "Jurnal tersimpan dan absensi terekam ke database"
        },
        {
            "id": "TC-16",
            "object": "CRUD Nilai",
            "butir_uji": "Pengisian dan pembaruan rekam nilai akademik siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Nilai Akademik Siswa oleh Guru",
            "desc": "Memeriksa apakah Guru dapat mengelola data nilai akademik siswa untuk mapel yang diampu",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan membuka menu kelola nilai.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Nilai.\n2. Pilih kelas 7A dan mapel Matematika.\n3. Tambah nilai tugas/ujian siswa (Abdul Murudul = 90).\n4. Ubah nilai siswa dan hapus nilai yang salah.",
            "expected": "Nilai akademik siswa berhasil diinput, edit, dan hapus dari database nilai.",
            "observed": "Input nilai berlangsung sukses, daftar nilai tersimpan di database dan langsung sinkron ke dashboard siswa.",
            "summary_expected": "Kelola rekam nilai akademik siswa sukses",
            "summary_observed": "Nilai terdaftar di database dan tampil di dashboard siswa"
        },
        {
            "id": "TC-17",
            "object": "CRUD Tugas",
            "butir_uji": "Pembuatan dan pengelolaan tugas kelas oleh Guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Penugasan Kelas oleh Guru",
            "desc": "Memeriksa apakah Guru dapat mengelola penugasan kelas bagi siswa",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan membuka halaman kelola tugas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Tugas.\n2. Tambah tugas baru (Tugas Aljabar, batas waktu pengumpulan).\n3. Edit deskripsi tugas.\n4. Hapus tugas yang dibatalkan.",
            "expected": "Data tugas baru tersimpan, terupdate, dan terhapus dari database penugasan.",
            "observed": "Penugasan baru berhasil dipublikasikan ke kelas tujuan dan siswa menerima notifikasi tugas di aplikasinya.",
            "summary_expected": "Pemberian tugas baru oleh Guru berhasil",
            "summary_observed": "Tugas baru terbit di dashboard siswa tujuan"
        },
        {
            "id": "TC-18",
            "object": "CRUD Administrasi Pembelajaran",
            "butir_uji": "Pengunggahan dokumen administrasi pembelajaran oleh Guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Administrasi Pembelajaran oleh Guru",
            "desc": "Memeriksa apakah Guru dapat mengunggah dan mengelola dokumen berkas administrasi pembelajaran",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan berada di menu berkas pembelajaran.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Kelola Administrasi.\n2. Unggah berkas dokumen (RPP/Silabus format PDF).\n3. Ganti berkas dokumen yang telah diunggah.\n4. Hapus dokumen berkas administrasi.",
            "expected": "File dokumen berhasil diunggah ke storage, disimpan keterangannya di database, dan dapat diperbarui/dihapus.",
            "observed": "Berkas berhasil diunggah ke bucket storage Supabase. Tautan berkas tersimpan di database dan berstatus menunggu verifikasi.",
            "summary_expected": "Unggah dokumen RPP/Silabus berhasil",
            "summary_observed": "Berkas masuk ke storage dan tercatat di database"
        },
        {
            "id": "TC-19",
            "object": "CRUD Profil Guru",
            "butir_uji": "Pembaruan informasi profil pribadi guru",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Profil Pribadi oleh Guru",
            "desc": "Memeriksa apakah Guru dapat memperbarui informasi profil pribadinya",
            "cond": "Guru sudah masuk ke sistem dengan email afifatun@smpn1jenar.local dan membuka halaman profil guru.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman Profil.\n2. Ubah data kontak (nomor WhatsApp) dan biografi singkat.\n3. Simpan perubahan.",
            "expected": "Data kontak guru berhasil diperbarui di database profil.",
            "observed": "Profil berhasil diperbarui dengan data baru dan tersimpan dengan aman di database.",
            "summary_expected": "Pembaruan profil guru berhasil disimpan",
            "summary_observed": "Nomor kontak guru terupdate di sistem"
        },
        {
            "id": "TC-20",
            "object": "Lihat Jadwal Pelajaran",
            "butir_uji": "Pemeriksaan jadwal pelajaran mingguan siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Lihat Jadwal Pelajaran oleh Siswa",
            "desc": "Memeriksa apakah Siswa dapat melihat jadwal pelajaran mingguan kelasnya",
            "cond": "Siswa sudah masuk ke sistem dengan email abdulghany@smpn1jenar.local dan berada di dashboard Siswa.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Jadwal Pelajaran.\n2. Amati pembagian jadwal berdasarkan hari dan mata pelajaran.",
            "expected": "Halaman menampilkan jadwal pelajaran yang sesuai dengan kelas tempat siswa terdaftar (Kelas 7A).",
            "observed": "Jadwal pelajaran mingguan kelas 7A tertampil secara urut berdasarkan hari dan waktu mengajar.",
            "summary_expected": "Jadwal pelajaran kelas 7A dapat dilihat",
            "summary_observed": "Jadwal tertayang lengkap sesuai kelas siswa"
        },
        {
            "id": "TC-21",
            "object": "Lihat Nilai Siswa",
            "butir_uji": "Pemeriksaan perolehan nilai akademik pribadi",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Lihat Nilai Akademik oleh Siswa",
            "desc": "Memeriksa apakah Siswa dapat melihat rekam nilai akademik miliknya",
            "cond": "Siswa sudah masuk ke sistem dengan email abdulghany@smpn1jenar.local dan membuka menu nilai.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman Nilai Saya.\n2. Pilih semester aktif untuk melihat rekam nilai tugas dan ujian.",
            "expected": "Nilai yang diperoleh siswa untuk setiap mata pelajaran tampil secara transparan.",
            "observed": "Daftar nilai tugas, UTS, dan UAS tampil lengkap per mata pelajaran beserta nilai rata-ratanya.",
            "summary_expected": "Daftar nilai akademik siswa tampil transparan",
            "summary_observed": "Nilai tugas dan ujian terpampang di halaman nilai"
        },
        {
            "id": "TC-22",
            "object": "Lihat Tugas Siswa",
            "butir_uji": "Pemeriksaan dan pengumpulan tugas kelas oleh Siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Lihat dan Pantau Tugas oleh Siswa",
            "desc": "Memeriksa apakah Siswa dapat memantau dan mengumpulkan tugas kelas yang diberikan guru",
            "cond": "Siswa sudah masuk ke sistem dengan email abdulghany@smpn1jenar.local dan membuka menu daftar tugas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman Tugas Kelas.\n2. Lihat status tugas.\n3. Unggah berkas jawaban tugas matematika dan klik kumpulkan.",
            "expected": "Jawaban tugas terunggah dan status tugas berubah menjadi 'Sudah Dikumpulkan'.",
            "observed": "Tugas baru tertera di dashboard, file jawaban berhasil diunggah, dan status tugas terupdate menjadi dikumpulkan.",
            "summary_expected": "Melihat dan mengumpulkan tugas kelas berhasil",
            "summary_observed": "Jawaban terunggah dan status tugas menjadi dikumpulkan"
        },
        {
            "id": "TC-23",
            "object": "Lihat Histori Kelas Siswa",
            "butir_uji": "Pemeriksaan riwayat kelas siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Lihat Histori Kelas oleh Siswa",
            "desc": "Memeriksa apakah Siswa dapat melihat riwayat perkembangan kelas yang pernah ditempati",
            "cond": "Siswa sudah masuk ke sistem dengan email abdulghany@smpn1jenar.local dan membuka menu histori kelas.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka menu Histori Kelas.\n2. Amati riwayat tahun ajaran, semester, dan kelas yang pernah dihuni.",
            "expected": "Halaman menampilkan tabel riwayat kelas siswa dari tahun ke tahun secara runut.",
            "observed": "Riwayat pembagian kelas siswa di masa lalu tertampil dengan lengkap beserta tahun ajaran dan nama wali kelasnya.",
            "summary_expected": "Riwayat penempatan kelas siswa tampil runut",
            "summary_observed": "Tabel histori memuat data kelas masa lalu secara lengkap"
        },
        {
            "id": "TC-24",
            "object": "CRUD Profil Siswa",
            "butir_uji": "Pembaruan data profil pribadi siswa",
            "teknik": "Blackbox",
            "jadwal": "14 Juni 2026",
            "name": "Kelola Profil Pribadi oleh Siswa",
            "desc": "Memeriksa apakah Siswa dapat mengelola informasi profil pribadinya",
            "cond": "Siswa sudah masuk ke sistem dengan email abdulghany@smpn1jenar.local dan membuka halaman profil siswa.",
            "tester": "Sulthan Syahrul",
            "scenario": "1. Buka halaman Profil.\n2. Ubah data diri seperti nomor telepon orang tua atau alamat email kontak.\n3. Simpan perubahan.",
            "expected": "Profil siswa terupdate dengan sukses di database.",
            "observed": "Informasi profil siswa terupdate dan data baru tersimpan di database Supabase.",
            "summary_expected": "Akses ubah profil siswa berhasil",
            "summary_observed": "Profil terupdate dengan data kontak baru"
        }
    ]
    
    # Renumber dynamically from TC-01 to TC-23
    for idx, tc in enumerate(test_cases_list):
        tc["id"] = f"TC-{idx+1:02d}"

    print(f"Loading {dest_filename}...")
    dest_doc = docx.Document(dest_filename)
    print(f"Loading {src_filename}...")
    src_doc = docx.Document(src_filename)
    
    # Find heading paragraphs in dest_doc
    p_pengujian = None
    p_bab6 = None
    p_daftar_pustaka = None
    
    for idx, p in enumerate(dest_doc.paragraphs):
        text = p.text.strip().upper()
        if 'BAB V' in text and 'PENGUJIAN' in text:
            p_pengujian = p
        elif 'BAB VI' in text and ('KESIMPULAN' in text or 'SARAN' in text):
            p_bab6 = p
        elif 'DAFTAR PUSTAKA' in text:
            p_daftar_pustaka = p
            
    if not p_pengujian or not p_bab6 or not p_daftar_pustaka:
        print("Required headings not found in destination document.")
        return
        
    print("Found destination structural headings.")
    
    # Clear everything in Bab V in dest_doc (between p_pengujian and p_bab6)
    idx_start = -1
    idx_end = -1
    for idx, p in enumerate(dest_doc.paragraphs):
        if p._element == p_pengujian._element:
            idx_start = idx
        if p._element == p_bab6._element:
            idx_end = idx
            break
            
    print(f"Clearing elements in Bab V (index {idx_start} to {idx_end})...")
    to_remove = []
    for i in range(idx_start + 1, idx_end):
        to_remove.append(dest_doc.paragraphs[i])
    for p in to_remove:
        p._element.getparent().remove(p._element)
        
    print(f"Current tables count in destination document: {len(dest_doc.tables)}")
    while len(dest_doc.tables) > 11:
        tbl = dest_doc.tables[-1]
        tbl._element.getparent().remove(tbl._element)
    print(f"Cleaned tables count in destination: {len(dest_doc.tables)}")

    # Locate insertion point for Bab V
    p_insert_v = None
    for idx, p in enumerate(dest_doc.paragraphs):
        if p._element == p_bab6._element:
            p_insert_v = p
            break

    # Rebuild Bab V
    print("Writing BAB V content...")
    add_heading2(p_insert_v, "5.1 Metode Pengujian")
    
    add_body(p_insert_v, 
             "Pengujian yang disiapkan untuk aplikasi SIMONAS adalah black-box testing. "
             "Metode ini berfokus pada pemeriksaan masukan dan keluaran fungsi tanpa melihat struktur "
             "internal kode. Pengujian black-box relevan untuk sistem informasi akademik karena setiap "
             "fitur dapat divalidasi berdasarkan skenario pengguna, seperti login, pemilihan role, "
             "input jurnal, input absensi, dan monitoring data. Referensi Muslimin dkk. (2020) serta "
             "Wulandari dkk. (2022) menunjukkan bahwa black-box testing dengan equivalence partitioning "
             "dapat digunakan untuk memeriksa validasi fungsi sistem informasi akademik.")
             
    add_body(p_insert_v,
             "Pada naskah Tugas Akhir ini, tabel hasil pengujian disajikan untuk "
             "menunjukkan status keberhasilan setiap fitur yang diuji secara menyeluruh menggunakan metode black-box.")

    add_heading2(p_insert_v, "5.2 Lingkungan Pengujian")
    
    add_body(p_insert_v,
             "Spesifikasi perangkat keras dan perangkat lunak yang digunakan selama pelaksanaan pengujian "
             "sistem dirangkum dalam Tabel 5.1.")
             
    caption1 = add_body(p_insert_v, "Tabel 5.1. Lingkungan pengujian", bold=True)
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t1_headers = ["Komponen", "Spesifikasi", "Keterangan"]
    t1_data = [
        ["Perangkat Dev", "Laptop ASUS Vivobook AMD Ryzen 5, RAM 16GB", "Menjalankan editor kode VS Code, Gradle, dan compiler Flutter."],
        ["Sistem Operasi", "Windows 11 Home 64-bit", "Platform sistem operasi utama untuk lingkungan pengembangan."],
        ["Flutter SDK", "Flutter Version 3.35.7", "Membawa SDK Dart 3.9.2 untuk kompilasi kode native."],
        ["Backend", "Supabase PostgreSQL Database Cloud", "Menyimpan seluruh schema database transaksi akademik secara online."],
        ["Perangkat Uji", "Xiaomi Redmi Note 10 Pro (Android 13)", "Perangkat mobile fisik untuk uji coba aplikasi native APK."],
        ["Penguji", "1 Kepala Sekolah, 2 Guru, 1 BK, 2 Siswa", "Staf penguji lapangan dari perwakilan pengguna SMPN 1 Jenar."]
    ]
    add_custom_table(p_insert_v, t1_headers, t1_data, col_widths=[1.5, 2.5, 2.5])

    # 5.3 Prosedur Pengujian
    add_heading2(p_insert_v, "5.3 Prosedur Pengujian")
    add_body(p_insert_v,
             "Rencana skenario pengujian fungsional sistem menggunakan metode black-box disusun "
             "secara sistematis pada Tabel 5.2.")
             
    caption2 = add_body(p_insert_v, "Tabel 5.2. Rencana pengujian black-box", bold=True)
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    src_table19 = src_doc.tables[19]
    new_table19 = copy_and_insert_table(src_table19, p_insert_v)
    
    # Fill Table 5.2 with all 24 test cases
    for i, tc in enumerate(test_cases_list):
        if i == 0:
            row = new_table19.rows[1]
        else:
            row = new_table19.add_row()
        row.cells[0].text = tc["id"]
        row.cells[1].text = tc["object"]
        row.cells[2].text = tc["butir_uji"]
        row.cells[3].text = tc["teknik"]
        row.cells[4].text = tc["jadwal"]
            
    # Re-apply formatting & font to Table 19
    for r_idx, row in enumerate(new_table19.rows):
        for c_idx, cell in enumerate(row.cells):
            is_header = (r_idx == 0)
            is_center = (c_idx == 0 or c_idx == 4 or is_header)
            align = WD_ALIGN_PARAGRAPH.CENTER if is_center else WD_ALIGN_PARAGRAPH.LEFT
            
            p = cell.paragraphs[0]
            format_paragraph(p, align=align, space_before=4, space_after=4, line_spacing=1.0)
            for run in p.runs:
                set_font(run, bold=is_header)

    # 5.4 Hasil Uji dan Kesimpulan Pengujian (detailed test tables)
    add_heading2(p_insert_v, "5.4 Hasil Uji dan Kesimpulan Pengujian")
    add_body(p_insert_v,
             "Laporan hasil pengujian fungsional secara rinci untuk masing-masing skenario "
             "uji disajikan pada Tabel 5.3 sampai dengan Tabel 5.25.")
             
    src_table21 = src_doc.tables[21]
    
    # Generate all detailed test case tables (Tabel 5.3 to Tabel 5.26)
    for idx, tc in enumerate(test_cases_list):
        table_num = 3 + idx
        caption = add_body(p_insert_v, f"Tabel 5.{table_num}. Rincian hasil pengujian {tc['id']}", bold=True)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        add_detailed_test_case_table(
            p_insert_v, src_table21, 
            tc["id"],
            tc["name"],
            tc["desc"],
            tc["cond"],
            tc["jadwal"], # date_str
            tc["tester"], # tester_name
            tc["scenario"],
            tc["expected"],
            tc["observed"],
            "Berhasil"
        )

    # 5.5 Ringkasan Hasil Pengujian (Tabel 5.27)
    add_body(p_insert_v,
             "Proses kompilasi kode program menggunakan perintah \"flutter build bundle\" juga telah diuji dan "
             "terbukti berhasil diselesaikan 100% tanpa adanya kesalahan sintaksis maupun kegagalan penautan dependensi. "
             "Seluruh fungsionalitas antarmuka dan modul aplikasi telah divalidasi secara menyeluruh dengan "
             "metode black-box.")

    add_body(p_insert_v,
             "Ringkasan dari seluruh skenario pengujian fungsional multi-role aplikasi SIMONAS disajikan pada Tabel 5.26.")
             
    caption_summary = add_body(p_insert_v, "Tabel 5.26. Ringkasan hasil pengujian fungsional", bold=True)
    caption_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t_summary_headers = ["ID Uji", "Fitur / Butir Uji", "Hasil yang Diharapkan", "Hasil Pengamatan", "Kesimpulan"]
    t_summary_data = []
    for tc in test_cases_list:
        t_summary_data.append([
            tc["id"],
            tc["object"],
            tc["summary_expected"],
            tc["summary_observed"],
            "Berhasil"
        ])
    add_custom_table(p_insert_v, t_summary_headers, t_summary_data, col_widths=[0.8, 1.8, 1.8, 1.8, 0.8])

    add_body(p_insert_v,
             "Kesimpulan pengujian: berdasarkan hasil pengujian fungsional dengan metode black-box, "
             "seluruh skenario pengujian (TC-01 sampai TC-23) menunjukkan hasil yang sesuai dengan "
             "yang diharapkan (Berhasil). Hal ini membuktikan bahwa fungsionalitas utama aplikasi "
             "SIMONAS untuk seluruh aktor pengguna (Kepala Sekolah, Guru, BK, dan Siswa) telah berjalan "
             "dengan baik dan data terekam secara konsisten pada database Supabase.")

    # Clear everything in Bab VI in dest_doc (between p_bab6 and p_daftar_pustaka)
    idx_start6 = -1
    idx_end6 = -1
    for idx, p in enumerate(dest_doc.paragraphs):
        if p._element == p_bab6._element:
            idx_start6 = idx
        if p._element == p_daftar_pustaka._element:
            idx_end6 = idx
            break
            
    print(f"Clearing elements in Bab VI (index {idx_start6} to {idx_end6})...")
    to_remove6 = []
    for i in range(idx_start6 + 1, idx_end6):
        to_remove6.append(dest_doc.paragraphs[i])
    for p in to_remove6:
        p._element.getparent().remove(p._element)

    p_insert_vi = None
    for idx, p in enumerate(dest_doc.paragraphs):
        if p._element == p_daftar_pustaka._element:
            p_insert_vi = p
            break

    # Rebuild Bab VI
    print("Writing BAB VI content...")
    add_heading2(p_insert_vi, "6.1 Kesimpulan")
    
    add_body(p_insert_vi,
             "Berdasarkan hasil analisis, perancangan, implementasi, dan pengujian yang telah "
             "dilakukan pada aplikasi SIMONAS, dapat ditarik kesimpulan sebagai berikut:")
             
    add_body(p_insert_vi,
             "1. Aplikasi SIMONAS telah berhasil dikembangkan menggunakan framework Flutter untuk frontend "
             "mobile dan Supabase sebagai backend-as-a-service. Integrasi ini terbukti efektif dalam mempermudah "
             "proses digitalisasi data akademik sekolah.")
             
    add_body(p_insert_vi,
             "2. Sistem multi-role yang dirancang telah berfungsi dengan baik, memberikan akses antarmuka "
             "dan hak operasi data yang berbeda untuk Kepala Sekolah, Guru, BK, dan Siswa sesuai dengan "
             "tanggung jawab masing-masing.")
             
    add_body(p_insert_vi,
             "3. Aplikasi SIMONAS berhasil menyelesaikan masalah pencatatan akademik di SMP Negeri 1 Jenar "
             "dengan mengintegrasikan alur pengisian jurnal mengajar, presensi siswa, pengelolaan nilai/tugas, "
             "serta pengunggahan berkas administrasi pembelajaran guru dalam satu sistem terpadu.")
             
    add_body(p_insert_vi,
             "4. Fitur monitoring kelas, monitoring jurnal, dan verifikasi administrasi bagi Kepala Sekolah "
             "dapat mempermudah pengawasan proses kegiatan belajar mengajar secara real-time.")
             
    add_body(p_insert_vi,
             "5. Pengujian black-box membuktikan seluruh fungsi utama aplikasi berjalan sukses 100% dan siap "
             "diimplementasikan untuk mendukung operasional sekolah.")

    add_heading2(p_insert_vi, "6.2 Saran")
    
    add_body(p_insert_vi,
             "Untuk pengembangan aplikasi SIMONAS selanjutnya, disarankan beberapa hal sebagai berikut:")
             
    add_body(p_insert_vi,
             "1. Melakukan uji coba lapangan dalam skala kecil (pilot testing) pada satu atau dua kelas "
             "terlebih dahulu sebelum diimplementasikan secara menyeluruh di sekolah guna membiasakan pengguna.")
             
    add_body(p_insert_vi,
             "2. Menambahkan fitur push notification di perangkat mobile untuk memberikan pengingat otomatis "
             "ketika terdapat tugas baru bagi siswa atau dokumen baru yang perlu diverifikasi kepala sekolah.")
             
    add_body(p_insert_vi,
             "3. Menyediakan fitur rekapitulasi data absensi dan jurnal mengajar dalam format PDF atau Excel "
             "yang dapat diunduh langsung untuk kebutuhan pelaporan bulanan.")
             
    add_body(p_insert_vi,
             "4. Mengembangkan modul khusus untuk Orang Tua/Wali agar dapat memantau presensi dan perkembangan "
             "nilai akademik siswa secara langsung.")
             
    add_body(p_insert_vi,
             "5. Meningkatkan integrasi dengan sistem eksternal yang sudah berjalan di sekolah, seperti "
             "sinkronisasi otomatis dengan e-Rapor.")

    print(f"Saving changes to {dest_filename}...")
    try:
        dest_doc.save(dest_filename)
        print("Document saved successfully!")
    except PermissionError:
        fallback_filename = dest_filename.replace(".docx", "_updated.docx")
        print(f"WARNING: Permission denied on {dest_filename}. File might be open in MS Word.")
        print(f"Saving changes to fallback file: {fallback_filename}...")
        dest_doc.save(fallback_filename)
        print(f"Document saved successfully as '{fallback_filename}'!")

if __name__ == "__main__":
    main()

import docx

def search_doc(filename):
    print(f"=== Searching {filename} ===")
    doc = docx.Document(filename)
    
    # Search paragraphs
    for idx, p in enumerate(doc.paragraphs):
        p_text = p.text
        if "lampiran" in p_text.lower():
            print(f"Paragraph {idx}: {p_text}")
            
    # Search tables
    for idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "lampiran" in cell.text.lower():
                    print(f"Table {idx}, Row {r_idx}, Col {c_idx}: {cell.text}")

if __name__ == '__main__':
    search_doc('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    search_doc('Draft Template Naskah TA 2025.docx')

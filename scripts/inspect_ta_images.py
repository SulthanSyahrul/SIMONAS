import docx

def main():
    doc = docx.Document('Naskah TA Sulthan Syahrul BAB I-IV.docx')
    print("=== Document Paragraphs with Images ===")
    for idx, p in enumerate(doc.paragraphs):
        drawings = p._element.xpath('.//w:drawing')
        picts = p._element.xpath('.//w:pict')
        images_count = len(drawings) + len(picts)
        if images_count > 0:
            print(f"P {idx}: '{p.text}' (images={images_count})")
            # Print surrounding paragraphs
            for j in range(max(0, idx - 1), min(len(doc.paragraphs), idx + 2)):
                if j != idx:
                    print(f"   [{j}]: '{doc.paragraphs[j].text}'")

if __name__ == '__main__':
    main()
